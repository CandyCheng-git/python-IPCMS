# report.py

from docx import Document
from docx.shared import Inches
from datetime import datetime
import os
from utils.utils import get_cur_location, get_sales_quantity
from utils.transformers import DataTransformer


class ReportGenerator:
    """Class to generate and manage reports."""

    def __init__(self, customer_manager, product_manager, order_manager, chart_generator):
        self.transform_data_method = DataTransformer()
        self.customer_manager = customer_manager
        self.product_manager = product_manager
        self.order_manager = order_manager
        self.chart_generator = chart_generator

    def create_report_content(self):
        """Generate the content of the report."""
        return {
            'kpis': self.gen_kpis_section(),
            'product_performance': self.gen_product_performance_section(),
            'order_customer_insights': self.gen_order_customer_insights_section(),
            'charts': self.get_chart_filenames(),
        }

    # Report Function gen_kpis_section: gen KPI's section
    def gen_kpis_section(self):
        tt_sales = sum(float(order.total_price) for order in self.order_manager.get_all_orders())
        tt_customers = len(self.customer_manager.get_all_customers())
        tt_orders = len(self.order_manager.get_all_orders())
        avg_order_value = tt_sales / tt_orders if tt_orders > 0 else 0
        customer_retention_rate = 100  # Placeholder logic

        return {
            'Total Sales': f"${tt_sales:,.2f}",
            'Total Customers': tt_customers,
            'Average Order Value': f"${avg_order_value:,.2f}",
            'Customer Retention Rate': f"{customer_retention_rate:.2f}%"  # Simplified
        }

    # Report Function gen_product_performance_section: gen product's performance section
    def gen_product_performance_section(self):
        pd_sales = {}
        for a_order in self.order_manager.get_all_orders():
            for op in a_order.order_products:
                if isinstance(op, dict):
                    pd_sales[op['product_name']] = pd_sales.get(op['product_name'], 0) + int(op['quantity'])
                elif hasattr(op, 'product_name') and hasattr(op, 'quantity'):
                    pd_sales[op.product_name] = pd_sales.get(op.product_name, 0) + int(op.quantity)
        best_sale = max(pd_sales.items(), key=get_sales_quantity, default=("N/A",))[0]
        the_low_stock = [f"{p.product_name} – {p.stock_quantity} units" for p in self.product_manager.get_all_products()
                         if p.stock_quantity < 10]

        return {
            'Best-Selling Product': best_sale,
            'Low Stock Alert': the_low_stock,
            'Product Table': [
                {'Product Name': p.product_name, 'Units Sold': pd_sales.get(p.product_name, 0),
                 'Total Revenue': f"${pd_sales.get(p.product_name, 0) * p.price:,.2f}"}
                for p in self.product_manager.get_all_products()
            ]
        }

    # Report Function gen_order_customer_insights_section: gen customer's insights section
    def gen_order_customer_insights_section(self):
        rc_orders = sorted(
            self.order_manager.get_all_orders(),
            key=self.transform_data_method.transform_order_date,
            reverse=True
        )[:5]
        cities_of_customer = {}
        for a_customer in self.customer_manager.get_all_customers():
            cities_of_customer[a_customer.city] = cities_of_customer.get(a_customer.city, 0) + 1

        return {
            'Recent Orders': [{'Order ID': o.order_id, 'Date': o.order_date, 'Total': f"${float(o.total_price):,.2f}"}
                              for o in rc_orders],
            'Customer Demographics': [{'City': city, 'Number of Customers': count} for city, count in
                                      cities_of_customer.items()]
        }

    @staticmethod
    def get_chart_filenames():
        """Get the filenames of the charts."""
        return [
            'top-selling_products.png',
            'customer_location_distribution.png',
            'customer_age_distribution.png'
        ]

    # Report Function create_report_document: gen report's Microsoft doc
    def create_report_document(self, ms_wd_doc, contents):
        # Add Title
        ms_wd_doc.add_heading('Marketing Report', 0)

        # 1st Section – KPIs
        ms_wd_doc.add_heading('1. KPIs', level=1)
        for key, value in contents['kpis'].items():
            ms_wd_doc.add_paragraph(f"• {key}: {value}")

        # 2nd Section – Product Performance
        ms_wd_doc.add_heading('2. Product Performance', level=1)
        ms_wd_doc.add_paragraph(f"• Best-Selling Product: {contents['product_performance']['Best-Selling Product']}")
        ms_wd_doc.add_paragraph(f"• Low Stock Alert:")
        for alert in contents['product_performance']['Low Stock Alert']:
            ms_wd_doc.add_paragraph(f"  - {alert}", style='List Bullet')

        # 3rd Section – Product Table
        ms_wd_doc.add_heading('3. Product Table', level=1)
        table = ms_wd_doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Product Name'
        hdr_cells[1].text = 'Units Sold'
        hdr_cells[2].text = 'Total Revenue'
        for a_item in contents['product_performance']['Product Table']:
            row_cells = table.add_row().cells
            row_cells[0].text = a_item['Product Name']
            row_cells[1].text = str(a_item['Units Sold'])
            row_cells[2].text = a_item['Total Revenue']

        # 4th Section – Order and Customer Insights
        ms_wd_doc.add_heading('4. Order and Customer Insights', level=1)
        table = ms_wd_doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Order ID'
        hdr_cells[1].text = 'Date'
        hdr_cells[2].text = 'Total'
        for a_item in contents['order_customer_insights']['Recent Orders']:
            row_cells = table.add_row().cells
            row_cells[0].text = a_item['Order ID']
            row_cells[1].text = a_item['Date']
            row_cells[2].text = a_item['Total']

        # 5th - Customer Demographics Table
        ms_wd_doc.add_heading('5. Customer Demographics', level=1)
        table = ms_wd_doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'City'
        hdr_cells[1].text = 'Number of Customers'
        for a_item in contents['order_customer_insights']['Customer Demographics']:
            row_cells = table.add_row().cells
            row_cells[0].text = a_item['City']
            row_cells[1].text = str(a_item['Number of Customers'])

        # 6th - Charts
        ms_wd_doc.add_heading('6. Charts', level=1)
        self.chart_generator.export_all_charts()
        for a_chart_filename in contents['charts']:
            if os.path.exists(a_chart_filename):
                ms_wd_doc.add_picture(a_chart_filename, width=Inches(6))
                ms_wd_doc.add_paragraph(a_chart_filename.replace('_', ' ').replace('.png', '').title())
            else:
                ms_wd_doc.add_paragraph(f"Chart '{a_chart_filename}' not found.")

    # Export Function export_report: Create a Report, then export
    def export_report(self):
        report_contents = self.create_report_content()
        set_filename = f"marketing_report_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        ms_word_doc = Document()
        self.create_report_document(ms_word_doc, report_contents)
        ms_word_doc.save(set_filename)
        print(f'Report "{set_filename}" created successfully. \nLocation: "{get_cur_location()}\\{set_filename}"')
