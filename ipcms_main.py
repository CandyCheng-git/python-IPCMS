
# Combined code for IPCMS (Integrated Project and Customer Management System)

""" Part 1 - Import Libraries """
from datetime import datetime
import matplotlib.pyplot as plt
import json
import numpy as np
import csv
import os
import re
from docx import Document
from docx.shared import Inches

""" Part 2 - Utility Functions """


# Class CheckValidator:: Group all validation functions
class CheckValidator:

    # Function is_numeric: Check input is an integer/number or not
    def is_numeric(self, b_check_value):
        if b_check_value is None or b_check_value == '':
            print("Error: Input cannot be empty.")
            return False
        try:
            int(b_check_value)
            return True
        except (Exception, ValueError, TypeError) as er_msg:
            print(f"Error: '{b_check_value}' is not a valid number. - {str(er_msg)}")
            return False

    # Function is_decimal: Check input is a float/decimal number or not
    def is_decimal(self, b_check_value):
        try:
            float(b_check_value)
            return True
        except (Exception, ValueError, TypeError) as er_msg:
            print(f"Error: '{b_check_value}' is not a valid decimal. - {str(er_msg)}")
            return False

    # Function is_alpha: Check input contains only alphabetic characters or not
    def is_alpha(self, b_check_value):
        if b_check_value is None or b_check_value == '':
            print("Error: Input cannot be empty.")
            return False
        if not b_check_value or not b_check_value.isalpha():
            print(f"Error: '{b_check_value}' contains non-alphabetic characters.")
            return False
        return True

    # Function is_valid_date_of_birth: Check input is a valid date of birth in YYYY-MM-DD format or not
    def is_valid_date_of_birth(self, b_check_value):
        if b_check_value is None or b_check_value == '':
            print("Error: Input cannot be empty.")
            return False
        parts_in_check_value = b_check_value.split('-')
        if len(parts_in_check_value) == 3 and all(self.is_numeric(part) for part in parts_in_check_value):
            year = int(parts_in_check_value[0])
            if 1900 <= year <= datetime.now().year:
                return True
        print(f"Error: '{b_check_value}' is not a valid date of birth.")
        return False

    # Function is_datetime: Check input is a valid datetime in YYYY-MM-DD HH:MM:SS format or not
    def is_datetime(self, b_check_value):
        key_pat = "%Y-%m-%d %H:%M:%S"
        try:
            datetime.strptime(b_check_value, key_pat)
            return True
        except (Exception, ValueError, TypeError) as er_msg:
            print(f"Error: '{b_check_value}' is not a valid datetime. - {str(er_msg)}")
            return False

    # Function is_phone_num: Check input is a valid Australian phone number (e.g.'0000-000-000') or not
    def is_phone_num(self, b_check_value):
        if b_check_value is None or b_check_value == '':
            print("Error: Input cannot be empty.")
            return False
        key_pat = r'^0\d{3}-\d{3}-\d{3}$'
        if re.match(key_pat, b_check_value):
            return True
        print(f"Error: '{b_check_value}' is not a valid Australian phone number. Suggest format : 0000-000-000 ")
        return False

    # Function is_email: Check an email address is valid or not
    def is_email(self, b_check_value):
        if b_check_value is None or b_check_value == '':
            print("Error: Input cannot be empty.")
            return False
        key_pat = r'^[\w\.-]+@[\w\.-]+\.\w+$'
        if re.match(key_pat, b_check_value):
            return True
        print(f"Error: '{b_check_value}' is not a valid email address. Suggest format : user@example.com ")
        return False

    # Function is_served_country: Check input contains a country or not
    def is_served_country(self, b_check_value, served_countries):
        if b_check_value is None or b_check_value == '':
            print("Error: Input cannot be empty.")
            return False
        if self.is_alpha(b_check_value) and b_check_value in served_countries:
            return True
        print(f"Error: '{b_check_value}' is not in the served countries.")
        return False

    # Function is_postcode: Check input is a valid postcode in the served country list or not
    def is_postcode(self, country_name, b_check_value, served_countries):
        if b_check_value is None or b_check_value == '':
            print("Error: Input cannot be empty.")
            return False
        key_pattern = served_countries.get(country_name, {}).get("Postcode Format")
        if key_pattern and re.match(key_pattern, b_check_value):
            return True
        print(f"Error: '{b_check_value}' is not a valid postcode for {country_name}.")
        return False


# Class DataTransformer:: Group all transform functions
class DataTransformer:
    # Helper function pad_string: Pad the string to match the target width
    def pad_string(self, input_string, target_width):
        input_string = str(input_string)
        return input_string + ' ' * max(target_width - len(input_string), 0)

    # Helper function to_std_dateformat: convert different date formats to YYYY-MM-DD for read_csv
    def to_std_dateformat(self, b_check_str_d):
        for fmt in ('%d/%m/%Y', '%d-%m-%Y', '%Y-%m-%d', '%Y/%m/%d'):
            try:
                return datetime.strptime(b_check_str_d, fmt).strftime('%Y-%m-%d')
            except ValueError:
                continue
        return None

    # Helper function to_std_datetimeformat: convert different date formats to YYYY-MM-DD HH:MM:SS for read_csv
    def to_std_datetimeformat(self, b_check_str_dt):
        for fmt in (
                '%d/%m/%Y %H:%M:%S', '%d-%m-%Y %H:%M:%S', '%Y-%m-%d %H:%M:%S', '%Y/%m/%d %H:%M:%S',
                '%d/%m/%Y %H:%M', '%d-%m-%Y %H:%M', '%Y-%m-%d %H:%M', '%Y/%m/%d %H:%M'
        ):
            try:
                return datetime.strptime(b_check_str_dt, fmt).strftime('%Y-%m-%d %H:%M:%S')
            except ValueError:
                continue
        print(f"Error: '{b_check_str_dt}' is not in a valid datetime format.")
        return None

    # Function to capitalize the first letter of each word - .title()
    def format_name(self, string):
        return string.title()


""" Part 3 - Class Objects Enterprise's Data """


# Class Customer:: A customer object
class Customer:
    # Constructor
    def __init__(
            self,
            first_name,
            last_name,
            dob,
            email,
            phone,
            country,
            city,
            postcode,
            created_at,
            updated_at,
    ):
        self.first_name = first_name
        self.last_name = last_name
        self.dob = dob
        self.email = email
        self.phone = phone
        self.country = country
        self.city = city
        self.postcode = postcode
        self.created_at = created_at
        self.updated_at = updated_at


# Class Product:: A product object
class Product:
    # Constructor
    def __init__(
            self,
            product_id,
            product_name,
            price,
            category,
            stock_quantity,
            created_at,
            updated_at,
    ):
        self.product_id = str(product_id)
        self.product_name = product_name
        self.price = float(price)
        self.category = category
        self.created_at = created_at
        self.updated_at = updated_at
        self.stock_quantity = int(stock_quantity) if CheckValidator().is_numeric(stock_quantity) else 0


# Class OrderedProduct:: An Ordered Product object
class OrderedProduct:
    # Constructor
    def __init__(
            self, product_id, product_name, quantity, price_per_unit, total_price
    ):
        self.product_id = str(product_id)
        self.product_name = product_name
        self.quantity = int(quantity)
        self.price_per_unit = float(price_per_unit)
        self.total_price = float(total_price)


# Class Order:: An Order object
class Order:
    # Constructor
    def __init__(
            self,
            order_id,
            order_date,
            order_status,
            total_price,
            order_products,
            payment_method,
            created_at,
            updated_at,
            customer_email,
    ):
        self.order_id = order_id
        self.order_date = order_date
        self.order_status = order_status
        self.total_price = float(total_price)
        self.payment_method = payment_method
        self.created_at = created_at
        self.updated_at = updated_at
        self.customer_email = customer_email
        self.order_products = [
            OrderedProduct(
                product_id=prd['product_id'],
                product_name=prd.get('product_name', ''),
                quantity=int(prd['quantity']),
                price_per_unit=float(prd['price_per_unit']),
                total_price=prd.get('total_price', int(prd['quantity']) * float(prd['price_per_unit']))
            ) if isinstance(prd, dict) else prd for prd in order_products
        ]


# Class CustomerController:: Manage all operations which is related to a Customer
class CustomerController:
    # Constructor
    def __init__(self, order_manager):
        self.customers = []
        self.check_valid_method = CheckValidator()
        self.order_manager = order_manager

    def create_customer(
            self,
            first_name,
            last_name,
            dob,
            email,
            phone,
            country,
            city,
            postcode,
            created_at,
            updated_at,
    ):
        new_customer = Customer(
            first_name,
            last_name,
            dob,
            email,
            phone,
            country,
            city,
            postcode,
            created_at,
            updated_at,
        )
        existing_customer = self.get_customer_by_email(new_customer.email)
        if existing_customer:
            self.update_customer_by_email(new_customer)
        else:
            self.customers.append(new_customer)

    def get_customer_by_email(self, tar_email):
        return next((customer for customer in self.customers if customer.email == tar_email), None)

    def update_customer_by_email(self, new_customer):
        existing_customer = self.get_customer_by_email(new_customer.email)
        if existing_customer:
            for idx, existing_customer in enumerate(self.customers):
                if existing_customer.email == new_customer.email:
                    new_customer.updated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    self.customers[idx] = new_customer
                    print(f"Success: Customer {new_customer.email} updated successfully.")
        else:
            print(f"Customer {new_customer.email} not found.")

    def delete_customer(self, email):
        existing_customer = self.get_customer_by_email(email)
        if existing_customer:
            confirmation = input(f"Are you sure you want to delete {email}? (y/n): ")
            if confirmation.lower() == 'y':
                # Assuming `self.order_manager` exists and manages orders related to the customer
                self.order_manager.delete_orders_by_customer_email(email)
                self.customers.remove(existing_customer)
                print(f"Customer {email} deleted.")
            else:
                print("Deletion cancelled.")
        else:
            print(f"Customer with email {email} does not exist.")

    def delete_orders_by_customer(self, email):
        existing_customer = self.get_customer_by_email(email)
        if existing_customer:
            self.order_manager.delete_orders_by_customer_email(email)
            self.customers.remove(existing_customer)
            print(f"Customer {email} deleted.")
            print(f"Customer with email {email} does not exist.")

    def get_all_customers(self):
        return self.customers


# Class ProductController:: Manage all operations which is related to a Product
class ProductController:
    # Constructor
    def __init__(self):
        self.products = []

    def create_product(
            self,
            product_id,
            product_name,
            price,
            category,
            stock_quantity,
            created_at,
            updated_at,
    ):
        if price < 0:
            print("Invalid price: Price cannot be negative.")
            return False
        product = Product(
            product_id,
            product_name,
            price,
            category,
            stock_quantity,
            created_at,
            updated_at,
        )
        existing_product = self.get_product_by_id(product_id)
        if existing_product:
            self.update_product_by_id(product)
        else:
            self.products.append(product)

    def update_product_by_id(self, new_product):
        if new_product is None:
            print("Error: Cannot update a None product.")
            return False

        for idx, existing_product in enumerate(self.products):
            if existing_product.product_id == new_product.product_id:
                existing_product.product_name = new_product.product_name
                existing_product.price = new_product.price
                existing_product.category = new_product.category
                existing_product.stock_quantity = new_product.stock_quantity
                existing_product.created_at = new_product.created_at
                existing_product.updated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                print(f"Product '{new_product.product_name}' updated successfully.")
                return True
        print(f"Product {new_product.product_id} not found.")
        return False

    def get_product_by_id(self, product_id):
        return next((prd for prd in self.products if prd.product_id == str(product_id).strip()), None)

    def get_all_products(self):
        return self.products


# Class OrderController:: Manage all operations which is related to an Order
class OrderController:
    # Constructor
    def __init__(self, product_manager):
        self.orders = []
        self.product_manager = product_manager
        self.items = []

    def create_order(
            self,
            order_id,
            order_date,
            total_price,
            order_products,
            payment_method,
            order_status,
            created_at,
            updated_at,
            customer_email
    ):
        for prd in order_products:
            if not self.product_manager.get_product_by_id(prd['product_id']):
                print(f"Product {prd['product_id']} does not exist. Order creation cancelled.")
                return None
        order = Order(
            order_id,
            order_date,
            order_status,
            total_price,
            order_products,
            payment_method,
            created_at,
            updated_at,
            customer_email
        )
        existing_order = self.get_order_by_id(order_id)
        if existing_order:
            self.update_order_by_id(order)
        else:
            self.orders.append(order)

    def update_order_by_id(self, new_order):
        for idx, existing_order in enumerate(self.orders):
            if existing_order.order_id == new_order.order_id:
                for key, value in vars(new_order).items():
                    setattr(existing_order, key, value)
                existing_order.updated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                return True
        print(f"Order {new_order.order_id} not found.")
        return False

    def delete_orders_by_customer_email(self, customer_email):
        self.orders = [
            order
            for order in self.orders
            if order.customer_email != customer_email
        ]

    def get_order_by_id(self, order_id):
        return next((order for order in self.orders if order.order_id == order_id), None)

    def get_all_orders(self):
        return self.orders


# Class EnterpriseData:: Load and Store initial data for the data from the shop
class EnterpriseData:
    # Constructor
    def __init__(self):
        self.col_widths = self.initi_col_widths()
        self.customers = self.load_customers()  # List to Objects
        self.products = self.load_products()  # List to Objects
        self.orders = self.load_orders()  # List to Objects
        self.employee_data_with_constraints = self.load_employee_data_with_constraints()

    def initi_col_widths(self):
        return {
            'customers': {
                "first_name": 12, "last_name": 12, "dob": 12, "email": 20,
                "phone": 12, "country": 12, "city": 12, "postcode": 4,
                "created_at": 19, "updated_at": 19,
            },
            'products': {
                "product_id": 3, "product_name": 30, "price": 10,
                "category": 10, "stock_quantity": 6, "created_at": 19,
                "updated_at": 19,
            },
            'orders': {
                "order_id": 6, "order_date": 10, "order_status": 10,
                "total_price": 10, "order_products": {
                    "product_id": 3, "product_name": 30, "quantity": 3,
                    "price_per_unit": 10, "total_price": 10,
                },
                "payment_method": 20, "customer_email": 20,
                "created_at": 19, "updated_at": 19,
            },
            'employee_login_dict': {
                "EmployeeID": 10,
                "Login Name": 8,
                "Password": 8
            },
            'currency_conversion_table': {
                "Country": 15,
                "Curr Code": 10,
                "Rate to AUD": 10
            },
            'employee_data_with_constraints': {
                "EmployeeID": 10,
                "Full Name": 20,
                "Department": 20,
                "Title": 20,
                "Base Yearly Salary": 15,
                "System Access Level": 5,
                "Workstation Name": 10,
                "Country": 15
            },
            'victoria_tax_table': {
                "Income Range Min": 20,
                "Income Range Max": 20,
                "Tax Rate (%)": 15
            },
            'payslip_table': {
                "ID": 20,
                "EmployeeID": 10,
                "Full Name": 20,
                "Department": 20,
                "Title": 20,
            }
        }

    def get_required_fields(self, category):
        if category in self.col_widths:
            return list(self.col_widths[category].keys())
        return []

    # Load data load_customers: all customers from customer_data
    def load_customers(self):
        customers = []
        for customer in self.customer_data():
            customers.append(Customer(
                first_name=customer['first_name'],
                last_name=customer['last_name'],
                dob=customer['dob'],
                email=customer['email'],
                phone=customer['phone'],
                country=customer['country'],
                city=customer['city'],
                postcode=customer['postcode'],
                created_at=customer['created_at'],
                updated_at=customer['updated_at']
            ))
        return customers

    # Load data load_products: all products from products_data
    def load_products(self):
        products = []
        for prd in self.products_data():
            products.append(Product(
                product_id=prd['product_id'],
                product_name=prd['product_name'],
                price=prd['price'],
                category=prd['category'],
                stock_quantity=prd['stock_quantity'],
                created_at=prd['created_at'],
                updated_at=prd['updated_at']
            ))
        return products

    # Load data load_orders: all products from orders_data
    def load_orders(self):
        orders = []
        for order in self.orders_data():
            orders.append(Order(
                order_id=order['order_id'],
                order_date=order['order_date'],
                total_price=order['total_price'],
                order_products=order['order_products'],
                payment_method=order['payment_method'],
                order_status=order['order_status'],
                created_at=order['created_at'],
                updated_at=order['updated_at'],
                customer_email=order['customer_email']
            ))
        return orders

    def load_employee_data_with_constraints(self):
        return [
            {"EmployeeID": 100000, "Full Name": "Admin User", "Department": "Administration", "Title": "Administrator", "Base Yearly Salary": 150000.00, "Country": "Australia", "System Access Level": "Admin", "Workstation Name": "WS-Admin"},
            {"EmployeeID": 100001, "Full Name": "John Doe", "Department": "IT Department", "Title": "System Administrator", "Base Yearly Salary": 80000.23, "Country": "Australia", "System Access Level": "Admin", "Workstation Name": "WS-1001"},
            {"EmployeeID": 100002, "Full Name": "Jane Smith", "Department": "HR Department", "Title": "HR Specialist", "Base Yearly Salary": 65000.75, "Country": "India", "System Access Level": "User", "Workstation Name": "WS-1002"},
            {"EmployeeID": 100003, "Full Name": "Albert Johnson", "Department": "Finance Department", "Title": "Accountant", "Base Yearly Salary": 90000.50, "Country": "China", "System Access Level": "User", "Workstation Name": "WS-1003"},
            {"EmployeeID": 100004, "Full Name": "Emily Brown", "Department": "Logistics Solution", "Title": "Flight Coordinator", "Base Yearly Salary": 78000.00, "Country": "Hong Kong", "System Access Level": "User", "Workstation Name": "WS-1004"},
            {"EmployeeID": 100005, "Full Name": "Michael Davis", "Department": "Sales Department", "Title": "Sales Executive", "Base Yearly Salary": 85000.00, "Country": "Australia", "System Access Level": "User", "Workstation Name": "WS-1005"},
            {"EmployeeID": 100006, "Full Name": "Sarah Wilson", "Department": "IT Department", "Title": "Developer", "Base Yearly Salary": 95000.00, "Country": "Vietnam", "System Access Level": "User", "Workstation Name": "WS-1006"},
            {"EmployeeID": 100007, "Full Name": "Robert Lee", "Department": "HR Department", "Title": "HR Manager", "Base Yearly Salary": 110000.00, "Country": "Malaysia", "System Access Level": "Admin", "Workstation Name": "WS-1007"},
            {"EmployeeID": 100008, "Full Name": "Jessica White", "Department": "Finance Department", "Title": "Payroll Specialist", "Base Yearly Salary": 67000.00, "Country": "India", "System Access Level": "User", "Workstation Name": "WS-1008"},
            {"EmployeeID": 100009, "Full Name": "Daniel Taylor", "Department": "Logistics Solution", "Title": "Shipping Coordinator", "Base Yearly Salary": 76000.00, "Country": "China", "System Access Level": "User", "Workstation Name": "WS-1009"},
            {"EmployeeID": 100010, "Full Name": "Amanda Thomas", "Department": "Sales Department", "Title": "Sales Associate", "Base Yearly Salary": 64000.00, "Country": "Australia", "System Access Level": "User", "Workstation Name": "WS-1010"},
            {"EmployeeID": 100011, "Full Name": "Andrew Martin", "Department": "IT Department", "Title": "IT Manager", "Base Yearly Salary": 120000.00, "Country": "Hong Kong", "System Access Level": "Admin", "Workstation Name": "WS-1011"},
            {"EmployeeID": 100012, "Full Name": "Megan Jackson", "Department": "HR Department", "Title": "HR Specialist", "Base Yearly Salary": 68000.00, "Country": "Vietnam", "System Access Level": "User", "Workstation Name": "WS-1012"},
            {"EmployeeID": 100013, "Full Name": "William Harris", "Department": "Finance Department", "Title": "Finance Manager", "Base Yearly Salary": 130000.00, "Country": "Malaysia", "System Access Level": "Admin", "Workstation Name": "WS-1013"},
            {"EmployeeID": 100014, "Full Name": "Olivia Robinson", "Department": "Logistics Solution", "Title": "Logistics Manager", "Base Yearly Salary": 115000.00, "Country": "China", "System Access Level": "Admin", "Workstation Name": "WS-1014"},
            {"EmployeeID": 100015, "Full Name": "James Clark", "Department": "Sales Department", "Title": "Sales Manager", "Base Yearly Salary": 125000.00, "Country": "Vietnam", "System Access Level": "Admin", "Workstation Name": "WS-1015"},
            {"EmployeeID": 100016, "Full Name": "Sophia Lewis", "Department": "IT Department", "Title": "Developer", "Base Yearly Salary": 85000.00, "Country": "India", "System Access Level": "User", "Workstation Name": "WS-1016"},
            {"EmployeeID": 100017, "Full Name": "Liam Walker", "Department": "HR Department", "Title": "HR Specialist", "Base Yearly Salary": 72000.00, "Country": "Australia", "System Access Level": "User", "Workstation Name": "WS-1017"},
            {"EmployeeID": 100018, "Full Name": "Chloe Hall", "Department": "Finance Department", "Title": "Payroll Specialist", "Base Yearly Salary": 66000.00, "Country": "Malaysia", "System Access Level": "User", "Workstation Name": "WS-1018"},
            {"EmployeeID": 100019, "Full Name": "Benjamin Allen", "Department": "Logistics Solution", "Title": "Shipping Coordinator", "Base Yearly Salary": 80000.00, "Country": "Hong Kong", "System Access Level": "User", "Workstation Name": "WS-1019"},
            {"EmployeeID": 100020, "Full Name": "Mia Young", "Department": "Sales Department", "Title": "Sales Executive", "Base Yearly Salary": 92000.00, "Country": "China", "System Access Level": "User", "Workstation Name": "WS-1020"}
        ]

    # staticmethod:: customer_data　(List)
    def customer_data(self):
        return [
            {
                'first_name': 'Matthew',
                'last_name': 'Contreras',
                'dob': '1995-10-07',
                'email': 'johnsmith@gmail.com',
                'phone': '0414-627-779',
                'country': 'Australia',
                'city': 'Melbourne',
                'postcode': '8018',
                'created_at': '2023-01-15 10:23:45',
                'updated_at': '2024-09-01 12:30:20',
            },
            {
                'first_name': 'Brenda',
                'last_name': 'Cannon',
                'dob': '1988-10-08',
                'email': 'janedoe@yahoo.com',
                'phone': '0495-195-938',
                'country': 'Australia',
                'city': 'Melbourne',
                'postcode': '6473',
                'created_at': '2023-03-10 14:45:50',
                'updated_at': '2024-09-02 09:15:12',
            },
            {
                'first_name': 'Kathryn',
                'last_name': 'Matthews',
                'dob': '1998-10-06',
                'email': 'bsmith@gmail.com',
                'phone': '0434-330-925',
                'country': 'Australia',
                'city': 'Melbourne',
                'postcode': '8552',
                'created_at': '2023-05-05 16:23:40',
                'updated_at': '2024-09-03 13:50:30',
            },
            {
                'first_name': 'Jeffrey',
                'last_name': 'Harris',
                'dob': '1997-10-06',
                'email': 'brobinson@duncan.com',
                'phone': '0432-125-147',
                'country': 'Australia',
                'city': 'Melbourne',
                'postcode': '6941',
                'created_at': '2023-07-18 11:12:56',
                'updated_at': '2024-09-04 10:24:42',
            },
            {
                'first_name': 'Megan',
                'last_name': 'Hunt',
                'dob': '1998-10-06',
                'email': 'sara11@gmail.com',
                'phone': '0437-725-970',
                'country': 'Australia',
                'city': 'Melbourne',
                'postcode': '9354',
                'created_at': '2023-09-10 15:35:18',
                'updated_at': '2024-09-05 14:45:05',
            },
            {
                'first_name': 'Michelle',
                'last_name': 'Perez',
                'dob': '2005-10-04',
                'email': 'lhall@todd.com',
                'phone': '0417-524-744',
                'country': 'Australia',
                'city': 'Melbourne',
                'postcode': '3806',
                'created_at': '2023-10-01 11:10:10',
                'updated_at': '2024-09-06 09:00:50',
            },
            {
                'first_name': 'Paul',
                'last_name': 'Wells',
                'dob': '1996-10-06',
                'email': 'tfowler@gmail.com',
                'phone': '0493-835-813',
                'country': 'Australia',
                'city': 'Melbourne',
                'postcode': '5999',
                'created_at': '2023-11-11 12:45:30',
                'updated_at': '2024-09-07 16:30:25',
            },
            {
                'first_name': 'April',
                'last_name': 'Jackson',
                'dob': '2003-10-05',
                'email': 'hmoore@gmail.com',
                'phone': '0467-509-349',
                'country': 'Australia',
                'city': 'Melbourne',
                'postcode': '3797',
                'created_at': '2023-12-02 17:22:45',
                'updated_at': '2024-09-08 10:45:00',
            },
            {
                'first_name': 'Steven',
                'last_name': 'Williams',
                'dob': '1996-10-06',
                'email': 'debkirby@gmail.com',
                'phone': '0420-868-981',
                'country': 'Australia',
                'city': 'Melbourne',
                'postcode': '4235',
                'created_at': '2023-05-20 09:13:29',
                'updated_at': '2024-09-09 11:25:35',
            },
            {
                'first_name': 'Johnny',
                'last_name': 'Hunter',
                'dob': '1972-10-12',
                'email': 'jonesh72@hotmail.com',
                'phone': '0416-324-280',
                'country': 'Australia',
                'city': 'Melbourne',
                'postcode': '5570',
                'created_at': '2024-01-15 14:00:00',
                'updated_at': '2024-09-10 08:32:17',
            },
            {
                'first_name': 'Kelly',
                'last_name': 'Nichols',
                'dob': '2000-10-05',
                'email': 'chris@hotmail.com',
                'phone': '0487-151-670',
                'country': 'Australia',
                'city': 'Melbourne',
                'postcode': '5978',
                'created_at': '2024-02-20 15:25:10',
                'updated_at': '2024-09-11 10:12:30',
            },
            {
                'first_name': 'Jeffrey',
                'last_name': 'Khan',
                'dob': '1978-10-11',
                'email': 'bradhunt@morris.com',
                'phone': '0431-567-236',
                'country': 'Australia',
                'city': 'Melbourne',
                'postcode': '2237',
                'created_at': '2024-03-01 18:12:45',
                'updated_at': '2024-09-12 09:45:00',
            },
            {
                'first_name': 'Kylie',
                'last_name': 'Nielsen',
                'dob': '1996-10-06',
                'email': 'susan95@gmail.com',
                'phone': '0478-672-709',
                'country': 'Australia',
                'city': 'Melbourne',
                'postcode': '5185',
                'created_at': '2024-04-15 13:50:20',
                'updated_at': '2024-09-13 08:40:10',
            },
            {
                'first_name': 'Benjamin',
                'last_name': 'Mclean',
                'dob': '1998-10-06',
                'email': 'dennis50@hotmail.com',
                'phone': '0457-727-934',
                'country': 'Australia',
                'city': 'Melbourne',
                'postcode': '4753',
                'created_at': '2024-05-18 14:25:45',
                'updated_at': '2024-09-14 11:15:35',
            },
            {
                'first_name': 'Larry',
                'last_name': 'Griffith',
                'dob': '2001-10-05',
                'email': 'pthompson@mitch.com',
                'phone': '0486-950-661',
                'country': 'Australia',
                'city': 'Melbourne',
                'postcode': '7479',
                'created_at': '2024-06-02 12:30:00',
                'updated_at': '2024-09-15 09:10:00',
            },
            {
                'first_name': 'Mathew',
                'last_name': 'Espinoza',
                'dob': '1995-10-07',
                'email': 'sabrinasmith@cross.bz',
                'phone': '0472-861-982',
                'country': 'Australia',
                'city': 'Melbourne',
                'postcode': '7563',
                'created_at': '2024-07-01 16:10:50',
                'updated_at': '2024-09-16 14:50:50',
            },
            {
                'first_name': 'Haley',
                'last_name': 'Horton',
                'dob': '1962-10-15',
                'email': 'pfitz@yahoo.com',
                'phone': '0477-365-869',
                'country': 'Australia',
                'city': 'Sydney',
                'postcode': '8538',
                'created_at': '2024-08-05 13:35:45',
                'updated_at': '2024-09-17 09:35:50',
            },
            {
                'first_name': 'Vincent',
                'last_name': 'Ferguson',
                'dob': '1975-10-12',
                'email': 'jessica11@hotmail.com',
                'phone': '0434-884-896',
                'country': 'Australia',
                'city': 'Brisbane',
                'postcode': '5312',
                'created_at': '2024-09-01 17:45:10',
                'updated_at': '2024-09-18 10:40:30',
            },
            {
                'first_name': 'Karen',
                'last_name': 'Lamb',
                'dob': '2006-10-04',
                'email': 'tonya67@gmail.com',
                'phone': '0487-741-344',
                'country': 'Australia',
                'city': 'Perth',
                'postcode': '7932',
                'created_at': '2024-10-01 18:15:35',
                'updated_at': '2024-09-19 12:25:55',
            },
            {
                'first_name': 'Matthew',
                'last_name': 'Jordan',
                'dob': '2005-10-04',
                'email': 'yvonne35@hotmail.com',
                'phone': '0447-179-830',
                'country': 'Australia',
                'city': 'Adelaide',
                'postcode': '3352',
                'created_at': '2024-11-12 13:12:45',
                'updated_at': '2024-09-20 16:45:30',
            },
        ]

    # staticmethod:: products_data　(List)
    def products_data(self):
        return [
            {
                "product_id": "001",
                "product_name": "EMO Robot (Lite)",
                "price": "311.98",
                "category": "Personal",
                "stock_quantity": "100",
                "created_at": "2024-01-10 09:15:23",
                "updated_at": "2024-09-01 12:45:56",
            },
            {
                "product_id": "002",
                "product_name": "EMO Robot (Classic)",
                "price": "389.98",
                "category": "Personal",
                "stock_quantity": "150",
                "created_at": "2024-02-11 10:45:30",
                "updated_at": "2024-09-02 14:35:40",
            },
            {
                "product_id": "003",
                "product_name": "EMO Robot (Dream)",
                "price": "436.78",
                "category": "Personal",
                "stock_quantity": "120",
                "created_at": "2024-03-12 14:30:00",
                "updated_at": "2024-09-03 16:25:42",
            },
            {
                "product_id": "004",
                "product_name": "EMO Robot (Ultimate)",
                "price": "514.78",
                "category": "Personal",
                "stock_quantity": "80",
                "created_at": "2024-04-14 08:12:45",
                "updated_at": "2024-09-04 09:20:50",
            },
            {
                "product_id": "005",
                "product_name": "Moxie Robot",
                "price": "1246.44",
                "category": "Educational",
                "stock_quantity": "50",
                "created_at": "2024-05-15 11:35:10",
                "updated_at": "2024-09-05 10:12:30",
            },
            {
                "product_id": "006",
                "product_name": "ROYBI Robot",
                "price": "310.44",
                "category": "Educational",
                "stock_quantity": "200",
                "created_at": "2024-06-16 10:15:22",
                "updated_at": "2024-09-06 12:10:15",
            },
            {
                "product_id": "007",
                "product_name": "Aibo Robot",
                "price": "4522.44",
                "category": "Pet",
                "stock_quantity": "30",
                "created_at": "2024-07-18 14:45:35",
                "updated_at": "2024-09-07 11:45:00",
            },
            {
                "product_id": "008",
                "product_name": "The Wild Robot - Roz",
                "price": "713.4",
                "category": "Love",
                "stock_quantity": "99",
                "created_at": "2024-09-18 08:45:35",
                "updated_at": "2024-09-30 11:45:00",
            },
        ]

    # staticmethod:: orders_data　(List)
    def orders_data(self):
        return [
            {
                "order_id": "PO0001",
                "order_date": "21-08-2024",
                "order_status": "Shipped",
                "total_price": "2043.50",
                "order_products": [
                    {
                        "product_id": "002",
                        "product_name": "EMO Robot (Classic)",
                        "quantity": "3",
                        "price_per_unit": "389.98",
                        "total_price": "1169.94",
                    },
                    {
                        "product_id": "003",
                        "product_name": "EMO Robot (Dream)",
                        "quantity": "2",
                        "price_per_unit": "436.78",
                        "total_price": "873.56",
                    },
                ],
                "payment_method": "Credit Card",
                "customer_email": "johnsmith@gmail.com",
                "created_at": "2024-08-21 09:00:00",
                "updated_at": "2024-09-01 12:00:00",
            },
            {
                "order_id": "PO0002",
                "order_date": "25-08-2024",
                "order_status": "Delivered",
                "total_price": "3472.50",
                "order_products": [
                    {
                        "product_id": "005",
                        "product_name": "Moxie Robot",
                        "quantity": "1",
                        "price_per_unit": "1246.44",
                        "total_price": "1246.44",
                    },
                    {
                        "product_id": "004",
                        "product_name": "EMO Robot (Ultimate)",
                        "quantity": "1",
                        "price_per_unit": "514.78",
                        "total_price": "514.78",
                    },
                    {
                        "product_id": "006",
                        "product_name": "ROYBI Robot",
                        "quantity": "3",
                        "price_per_unit": "310.44",
                        "total_price": "931.32",
                    },
                    {
                        "product_id": "002",
                        "product_name": "EMO Robot (Classic)",
                        "quantity": "2",
                        "price_per_unit": "389.98",
                        "total_price": "779.96",
                    },
                ],
                "payment_method": "PayPal",
                "customer_email": "janedoe@yahoo.com",
                "created_at": "2024-08-25 10:30:00",
                "updated_at": "2024-09-02 15:00:00",
            },
            {
                "order_id": "PO0003",
                "order_date": "16-08-2024",
                "order_status": "Shipped",
                "total_price": "5458.38",
                "order_products": [
                    {
                        "product_id": "007",
                        "product_name": "Aibo Robot",
                        "quantity": "1",
                        "price_per_unit": "4522.44",
                        "total_price": "4522.44",
                    },
                    {
                        "product_id": "001",
                        "product_name": "EMO Robot (Lite)",
                        "quantity": "3",
                        "price_per_unit": "311.98",
                        "total_price": "935.94",
                    },
                ],
                "payment_method": "Bank Transfer",
                "customer_email": "bjohnson@hotmail.com",
                "created_at": "2024-08-16 09:30:00",
                "updated_at": "2024-09-03 14:30:00",
            },
            {
                "order_id": "PO0004",
                "order_date": "09-08-2024",
                "order_status": "Delivered",
                "total_price": "14971.24",
                "order_products": [
                    {
                        "product_id": "007",
                        "product_name": "Aibo Robot",
                        "quantity": "3",
                        "price_per_unit": "4522.44",
                        "total_price": "13567.32",
                    },
                    {
                        "product_id": "002",
                        "product_name": "EMO Robot (Classic)",
                        "quantity": "2",
                        "price_per_unit": "389.98",
                        "total_price": "779.96",
                    },
                    {
                        "product_id": "001",
                        "product_name": "EMO Robot (Lite)",
                        "quantity": "2",
                        "price_per_unit": "311.98",
                        "total_price": "623.96",
                    },
                ],
                "payment_method": "Credit Card",
                "customer_email": "brobinson@duncan-wal",
                "created_at": "2024-08-09 11:45:00",
                "updated_at": "2024-09-04 16:00:00",
            },
            {
                "order_id": "PO0005",
                "order_date": "30-08-2024",
                "order_status": "Pending",
                "total_price": "9712.48",
                "order_products": [
                    {
                        "product_id": "005",
                        "product_name": "Moxie Robot",
                        "quantity": "3",
                        "price_per_unit": "1246.44",
                        "total_price": "3739.32",
                    },
                    {
                        "product_id": "001",
                        "product_name": "EMO Robot (Lite)",
                        "quantity": "2",
                        "price_per_unit": "311.98",
                        "total_price": "623.96",
                    },
                    {
                        "product_id": "003",
                        "product_name": "EMO Robot (Dream)",
                        "quantity": "1",
                        "price_per_unit": "436.78",
                        "total_price": "436.78",
                    },
                    {
                        "product_id": "007",
                        "product_name": "Aibo Robot",
                        "quantity": "1",
                        "price_per_unit": "4522.44",
                        "total_price": "4522.44",
                    },
                    {
                        "product_id": "002",
                        "product_name": "EMO Robot (Classic)",
                        "quantity": "1",
                        "price_per_unit": "389.98",
                        "total_price": "389.98",
                    },
                ],
                "payment_method": "Credit Card",
                "customer_email": "sara11@ward.com",
                "created_at": "2024-08-30 08:15:00",
                "updated_at": "2024-09-05 09:10:00",
            },
            {
                "order_id": "PO0006",
                "order_date": "02-09-2024",
                "order_status": "Shipped",
                "total_price": "311.98",
                "order_products": [
                    {
                        "product_id": "001",
                        "product_name": "EMO Robot (Lite)",
                        "quantity": "1",
                        "price_per_unit": "311.98",
                        "total_price": "311.98",
                    }
                ],
                "payment_method": "PayPal",
                "customer_email": "lhall@todd.com",
                "created_at": "2024-09-02 10:00:00",
                "updated_at": "2024-09-06 12:00:00",
            },
            {
                "order_id": "PO0007",
                "order_date": "05-09-2024",
                "order_status": "Delivered",
                "total_price": "1334.28",
                "order_products": [
                    {
                        "product_id": "006",
                        "product_name": "ROYBI Robot",
                        "quantity": "2",
                        "price_per_unit": "310.44",
                        "total_price": "620.88",
                    },
                    {
                        "product_id": "008",
                        "product_name": "The Wild Robot - Roz",
                        "quantity": "1",
                        "price_per_unit": "713.40",
                        "total_price": "713.40",
                    },
                ],
                "payment_method": "Credit Card",
                "customer_email": "tfowler@gmail.com",
                "created_at": "2024-09-05 11:30:00",
                "updated_at": "2024-09-07 14:20:00",
            },
            {
                "order_id": "PO0008",
                "order_date": "08-09-2024",
                "order_status": "Pending",
                "total_price": "779.96",
                "order_products": [
                    {
                        "product_id": "002",
                        "product_name": "EMO Robot (Classic)",
                        "quantity": "2",
                        "price_per_unit": "389.98",
                        "total_price": "779.96",
                    }
                ],
                "payment_method": "Bank Transfer",
                "customer_email": "hmoore@gmail.com",
                "created_at": "2024-09-08 09:15:00",
                "updated_at": "2024-09-08 09:15:00",
            },
            {
                "order_id": "PO0009",
                "order_date": "10-09-2024",
                "order_status": "Shipped",
                "total_price": "436.78",
                "order_products": [
                    {
                        "product_id": "003",
                        "product_name": "EMO Robot (Dream)",
                        "quantity": "1",
                        "price_per_unit": "436.78",
                        "total_price": "436.78",
                    }
                ],
                "payment_method": "Credit Card",
                "customer_email": "debbie75@kirby.com",
                "created_at": "2024-09-10 14:45:00",
                "updated_at": "2024-09-09 11:25:35",
            },
            {
                "order_id": "PO0010",
                "order_date": "12-09-2024",
                "order_status": "Delivered",
                "total_price": "310.44",
                "order_products": [
                    {
                        "product_id": "006",
                        "product_name": "ROYBI Robot",
                        "quantity": "1",
                        "price_per_unit": "310.44",
                        "total_price": "310.44",
                    }
                ],
                "payment_method": "PayPal",
                "customer_email": "jonesh@hotmail.com",
                "created_at": "2024-09-12 16:00:00",
                "updated_at": "2024-09-10 08:32:17",
            },
            {
                "order_id": "PO0011",
                "order_date": "14-09-2024",
                "order_status": "Pending",
                "total_price": "514.78",
                "order_products": [
                    {
                        "product_id": "004",
                        "product_name": "EMO Robot (Ultimate)",
                        "quantity": "1",
                        "price_per_unit": "514.78",
                        "total_price": "514.78",
                    }
                ],
                "payment_method": "Credit Card",
                "customer_email": "chrisg14@hotmail.com",
                "created_at": "2024-09-14 10:00:00",
                "updated_at": "2024-09-11 10:12:30",
            },
            {
                "order_id": "PO0012",
                "order_date": "16-09-2024",
                "order_status": "Shipped",
                "total_price": "935.94",
                "order_products": [
                    {
                        "product_id": "001",
                        "product_name": "EMO Robot (Lite)",
                        "quantity": "3",
                        "price_per_unit": "311.98",
                        "total_price": "935.94",
                    }
                ],
                "payment_method": "PayPal",
                "customer_email": "bradleyhunt@morris-s",
                "created_at": "2024-09-16 12:30:00",
                "updated_at": "2024-09-12 09:45:00",
            },
            {
                "order_id": "PO0013",
                "order_date": "18-09-2024",
                "order_status": "Delivered",
                "total_price": "1246.44",
                "order_products": [
                    {
                        "product_id": "005",
                        "product_name": "Moxie Robot",
                        "quantity": "1",
                        "price_per_unit": "1246.44",
                        "total_price": "1246.44",
                    }
                ],
                "payment_method": "Bank Transfer",
                "customer_email": "susan95@horton.com",
                "created_at": "2024-09-18 15:45:00",
                "updated_at": "2024-09-13 08:40:10",
            },
            {
                "order_id": "PO0014",
                "order_date": "20-09-2024",
                "order_status": "Shipped",
                "total_price": "623.96",
                "order_products": [
                    {
                        "product_id": "001",
                        "product_name": "EMO Robot (Lite)",
                        "quantity": "2",
                        "price_per_unit": "311.98",
                        "total_price": "623.96",
                    }
                ],
                "payment_method": "Credit Card",
                "customer_email": "dennis50@hotmail.com",
                "created_at": "2024-09-20 14:20:00",
                "updated_at": "2024-09-14 11:15:35",
            },
            {
                "order_id": "PO0015",
                "order_date": "22-09-2024",
                "order_status": "Delivered",
                "total_price": "1472.14",
                "order_products": [
                    {
                        "product_id": "008",
                        "product_name": "The Wild Robot - Roz",
                        "quantity": "2",
                        "price_per_unit": "713.40",
                        "total_price": "1426.80",
                    },
                    {
                        "product_id": "002",
                        "product_name": "EMO Robot (Classic)",
                        "quantity": "1",
                        "price_per_unit": "389.98",
                        "total_price": "389.98",
                    },
                ],
                "payment_method": "PayPal",
                "customer_email": "pthompson@mitchell-w",
                "created_at": "2024-09-22 11:00:00",
                "updated_at": "2024-09-15 09:10:00",
            },
            {
                "order_id": "PO0016",
                "order_date": "24-09-2024",
                "order_status": "Pending",
                "total_price": "4522.44",
                "order_products": [
                    {
                        "product_id": "007",
                        "product_name": "Aibo Robot",
                        "quantity": "1",
                        "price_per_unit": "4522.44",
                        "total_price": "4522.44",
                    }
                ],
                "payment_method": "Credit Card",
                "customer_email": "sabrinasmith@cross.b",
                "created_at": "2024-09-24 13:45:00",
                "updated_at": "2024-09-16 14:50:50",
            },
            {
                "order_id": "PO0017",
                "order_date": "26-09-2024",
                "order_status": "Shipped",
                "total_price": "389.98",
                "order_products": [
                    {
                        "product_id": "002",
                        "product_name": "EMO Robot (Classic)",
                        "quantity": "1",
                        "price_per_unit": "389.98",
                        "total_price": "389.98",
                    }
                ],
                "payment_method": "PayPal",
                "customer_email": "pfitzgerald@yahoo.co",
                "created_at": "2024-09-26 15:30:00",
                "updated_at": "2024-09-17 09:35:50",
            },
            {
                "order_id": "PO0018",
                "order_date": "28-09-2024",
                "order_status": "Delivered",
                "total_price": "310.44",
                "order_products": [
                    {
                        "product_id": "006",
                        "product_name": "ROYBI Robot",
                        "quantity": "1",
                        "price_per_unit": "310.44",
                        "total_price": "310.44",
                    }
                ],
                "payment_method": "Bank Transfer",
                "customer_email": "jessica11@hotmail.co",
                "created_at": "2024-09-28 09:50:00",
                "updated_at": "2024-09-18 10:40:30",
            },
            {
                "order_id": "PO0019",
                "order_date": "30-09-2024",
                "order_status": "Pending",
                "total_price": "847.22",
                "order_products": [
                    {
                        "product_id": "003",
                        "product_name": "EMO Robot (Dream)",
                        "quantity": "1",
                        "price_per_unit": "436.78",
                        "total_price": "436.78",
                    },
                    {
                        "product_id": "001",
                        "product_name": "EMO Robot (Lite)",
                        "quantity": "1",
                        "price_per_unit": "311.98",
                        "total_price": "311.98",
                    },
                    {
                        "product_id": "006",
                        "product_name": "ROYBI Robot",
                        "quantity": "1",
                        "price_per_unit": "310.44",
                        "total_price": "310.44",
                    },
                ],
                "payment_method": "Credit Card",
                "customer_email": "tonya67@durham.com",
                "created_at": "2024-09-30 11:15:00",
                "updated_at": "2024-09-19 12:25:55",
            },
            {
                "order_id": "PO0020",
                "order_date": "02-10-2024",
                "order_status": "Shipped",
                "total_price": "311.98",
                "order_products": [
                    {
                        "product_id": "001",
                        "product_name": "EMO Robot (Lite)",
                        "quantity": "1",
                        "price_per_unit": "311.98",
                        "total_price": "311.98",
                    }
                ],
                "payment_method": "PayPal",
                "customer_email": "yvonne35@hotmail.com",
                "created_at": "2024-10-02 13:00:00",
                "updated_at": "2024-09-20 16:45:30",
            },
            {
                "order_id": "PO0021",
                "order_date": "04-10-2024",
                "order_status": "Delivered",
                "total_price": "1521.52",
                "order_products": [
                    {
                        "product_id": "004",
                        "product_name": "EMO Robot (Ultimate)",
                        "quantity": "2",
                        "price_per_unit": "514.78",
                        "total_price": "1029.56",
                    },
                    {
                        "product_id": "002",
                        "product_name": "EMO Robot (Classic)",
                        "quantity": "1",
                        "price_per_unit": "389.98",
                        "total_price": "389.98",
                    },
                    {
                        "product_id": "006",
                        "product_name": "ROYBI Robot",
                        "quantity": "1",
                        "price_per_unit": "310.44",
                        "total_price": "310.44",
                    },
                ],
                "payment_method": "Credit Card",
                "customer_email": "ystevenson@gmail.com",
                "created_at": "2024-10-04 14:45:00",
                "updated_at": "2024-09-21 15:25:45",
            },
            {
                "order_id": "PO0022",
                "order_date": "06-10-2024",
                "order_status": "Pending",
                "total_price": "713.40",
                "order_products": [
                    {
                        "product_id": "008",
                        "product_name": "The Wild Robot - Roz",
                        "quantity": "1",
                        "price_per_unit": "713.40",
                        "total_price": "713.40",
                    }
                ],
                "payment_method": "PayPal",
                "customer_email": "sarahnguyen@miranda-",
                "created_at": "2024-10-06 16:30:00",
                "updated_at": "2024-09-22 14:50:30",
            },
            {
                "order_id": "PO0023",
                "order_date": "08-10-2024",
                "order_status": "Shipped",
                "total_price": "623.96",
                "order_products": [
                    {
                        "product_id": "001",
                        "product_name": "EMO Robot (Lite)",
                        "quantity": "2",
                        "price_per_unit": "311.98",
                        "total_price": "623.96",
                    }
                ],
                "payment_method": "Credit Card",
                "customer_email": "david56@mclaughlin.c",
                "created_at": "2024-10-08 09:00:00",
                "updated_at": "2024-09-23 10:12:20",
            },
            {
                "order_id": "PO0024",
                "order_date": "10-10-2024",
                "order_status": "Delivered",
                "total_price": "389.98",
                "order_products": [
                    {
                        "product_id": "002",
                        "product_name": "EMO Robot (Classic)",
                        "quantity": "1",
                        "price_per_unit": "389.98",
                        "total_price": "389.98",
                    }
                ],
                "payment_method": "PayPal",
                "customer_email": "zavalamason@hotmail.",
                "created_at": "2024-10-10 11:15:00",
                "updated_at": "2024-09-24 12:35:00",
            },
            {
                "order_id": "PO0025",
                "order_date": "12-10-2024",
                "order_status": "Pending",
                "total_price": "1246.44",
                "order_products": [
                    {
                        "product_id": "005",
                        "product_name": "Moxie Robot",
                        "quantity": "1",
                        "price_per_unit": "1246.44",
                        "total_price": "1246.44",
                    }
                ],
                "payment_method": "Credit Card",
                "customer_email": "phyllis33@hotmail.co",
                "created_at": "2024-10-12 13:30:00",
                "updated_at": "2024-09-25 11:45:30",
            },
            {
                "order_id": "PO0026",
                "order_date": "14-10-2024",
                "order_status": "Shipped",
                "total_price": "2368.62",
                "order_products": [
                    {
                        "product_id": "005",
                        "product_name": "Moxie Robot",
                        "quantity": "1",
                        "price_per_unit": "1246.44",
                        "total_price": "1246.44",
                    },
                    {
                        "product_id": "003",
                        "product_name": "EMO Robot (Dream)",
                        "quantity": "2",
                        "price_per_unit": "436.78",
                        "total_price": "873.56",
                    },
                    {
                        "product_id": "006",
                        "product_name": "ROYBI Robot",
                        "quantity": "1",
                        "price_per_unit": "310.44",
                        "total_price": "310.44",
                    },
                ],
                "payment_method": "PayPal",
                "customer_email": "phancock@hill-delgad",
                "created_at": "2024-10-14 15:45:00",
                "updated_at": "2024-09-26 13:15:00",
            },
            {
                "order_id": "PO0027",
                "order_date": "16-10-2024",
                "order_status": "Delivered",
                "total_price": "2335.94",
                "order_products": [
                    {
                        "product_id": "007",
                        "product_name": "Aibo Robot",
                        "quantity": "1",
                        "price_per_unit": "4522.44",
                        "total_price": "2261.22",
                    },
                    {
                        "product_id": "001",
                        "product_name": "EMO Robot (Lite)",
                        "quantity": "3",
                        "price_per_unit": "311.98",
                        "total_price": "935.94",
                    },
                ],
                "payment_method": "Bank Transfer",
                "customer_email": "bhess@hawkins-herrer",
                "created_at": "2024-10-16 17:00:00",
                "updated_at": "2024-09-27 14:25:10",
            },
            {
                "order_id": "PO0028",
                "order_date": "18-10-2024",
                "order_status": "Pending",
                "total_price": "826.76",
                "order_products": [
                    {
                        "product_id": "004",
                        "product_name": "EMO Robot (Ultimate)",
                        "quantity": "1",
                        "price_per_unit": "514.78",
                        "total_price": "514.78",
                    },
                    {
                        "product_id": "006",
                        "product_name": "ROYBI Robot",
                        "quantity": "1",
                        "price_per_unit": "310.44",
                        "total_price": "310.44",
                    },
                ],
                "payment_method": "Credit Card",
                "customer_email": "brandon84@ferguson.c",
                "created_at": "2024-10-18 09:15:00",
                "updated_at": "2024-09-28 10:55:30",
            },
            {
                "order_id": "PO0029",
                "order_date": "20-10-2024",
                "order_status": "Shipped",
                "total_price": "1646.16",
                "order_products": [
                    {
                        "product_id": "006",
                        "product_name": "ROYBI Robot",
                        "quantity": "4",
                        "price_per_unit": "310.44",
                        "total_price": "1241.76",
                    },
                    {
                        "product_id": "002",
                        "product_name": "EMO Robot (Classic)",
                        "quantity": "1",
                        "price_per_unit": "389.98",
                        "total_price": "389.98",
                    },
                ],
                "payment_method": "PayPal",
                "customer_email": "anne89@barnes.com",
                "created_at": "2024-10-20 11:30:00",
                "updated_at": "2024-09-29 11:30:15",
            },
            {
                "order_id": "PO0030",
                "order_date": "22-10-2024",
                "order_status": "Delivered",
                "total_price": "1646.16",
                "order_products": [
                    {
                        "product_id": "003",
                        "product_name": "EMO Robot (Dream)",
                        "quantity": "2",
                        "price_per_unit": "436.78",
                        "total_price": "873.56",
                    },
                    {
                        "product_id": "008",
                        "product_name": "The Wild Robot - Roz",
                        "quantity": "1",
                        "price_per_unit": "713.40",
                        "total_price": "713.40",
                    },
                ],
                "payment_method": "Bank Transfer",
                "customer_email": "wdavis@hotmail.com",
                "created_at": "2024-10-22 13:45:00",
                "updated_at": "2024-09-30 09:20:50",
            },
        ]




""" Part 4 - Class Objects ERP System's Data """


# Class ERPSystemData:: Load and Store initial data which is only for the ERP system
class ERPSystemData:
    # Constructor
    def __init__(self):
        self.welcome_message = self.initi_welcome_message()
        self.full_served_countries = self.load_full_served_countries()
        self.enterprise_data = EnterpriseData()
        self.custom_headers = {
            'customers': self.gen_customer_headers(),
            'products': self.gen_product_headers(),
            'orders': self.gen_order_headers(),
        }
        self.payslip_table = self.load_payslip_table()
        self.employee_login_dict = self.load_employee_login_dict()
        self.currency_conversion_table = self.load_currency_conversion_table()
        self.victoria_tax_table = self.load_victoria_tax_table()

    # Load data load_full_served_countries: Combine served_countries and postcode_format
    def load_full_served_countries(self):
        full_served_countries = {}
        served_countries = self.served_countries()
        served_postcode_format = self.served_postcode_format()

        for a_country in served_countries:
            country_name = a_country["Country"]
            full_served_countries[country_name] = {
                "Country Code": a_country["Country Code"],
                "Curr Code": a_country["Curr Code"],
                "Rate to AUD": a_country["Rate to AUD"]
            }

        for a_country_name, the_postcode_format in served_postcode_format:
            if a_country_name in full_served_countries:
                full_served_countries[a_country_name]["Postcode Format"] = the_postcode_format

        return full_served_countries

    def load_payslip_table(self):
        return list(())

    # Employee Login Table Data (Dictionary)
    def load_employee_login_dict(self):
        return {
            100000: {"Login Name": "admin", "Password": "admin123"},
            100001: {"Login Name": "jdoe", "Password": "e4a5Tc2m"},
            100002: {"Login Name": "jsmith", "Password": "9TpXk8Mb"},
            100003: {"Login Name": "ajohnson", "Password": "Ls7qP5Bm"},
            100004: {"Login Name": "ebrown", "Password": "Tn4kW2Vh"},
            100005: {"Login Name": "mdavis", "Password": "6QwEr3Xj"},
            100006: {"Login Name": "swilson", "Password": "Jh9qL6Pn"},
            100007: {"Login Name": "rlee", "Password": "Ap8sK2Qw"},
            100008: {"Login Name": "jwhite", "Password": "Dm3nP9Xh"},
            100009: {"Login Name": "dtaylor", "Password": "Fs6vL4Xn"},
            100010: {"Login Name": "athomas", "Password": "Zk2sW8Tj"},
            100011: {"Login Name": "amartin", "Password": "Xr3vP9Qn"},
            100012: {"Login Name": "mjackson", "Password": "Yn7mK4Ws"},
            100013: {"Login Name": "wharris", "Password": "Qj8tL6Pn"},
            100014: {"Login Name": "orobinson", "Password": "Vs4mP7Xn"},
            100015: {"Login Name": "jclark", "Password": "Hs9vK3Pm"},
            100016: {"Login Name": "slewis", "Password": "Bm7pW2Qx"},
            100017: {"Login Name": "lwalker", "Password": "Kn5qV9Xs"},
            100018: {"Login Name": "chall", "Password": "Ps6vL4Xj"},
            100019: {"Login Name": "ballen", "Password": "Zk3mW8Vn"},
            100020: {"Login Name": "myoung", "Password": "Qs7tP4Wm"}
        }

    # Currency Conversion Table Data (Tuple)
    def load_currency_conversion_table(self):
        return (
            {"Country": "Australia", "Curr Code": "AUD", "Rate to AUD": 1.0},
            {"Country": "Hong Kong", "Curr Code": "HKD", "Rate to AUD": 5.6},
            {"Country": "China", "Curr Code": "CNY", "Rate to AUD": 4.6},
            {"Country": "Malaysia", "Curr Code": "MYR", "Rate to AUD": 3.2},
            {"Country": "Vietnam", "Curr Code": "VND", "Rate to AUD": 16000.0},
            {"Country": "India", "Curr Code": "INR", "Rate to AUD": 55.3}
        )

    # Victoria Tax Table Data (Set)
    def load_victoria_tax_table(self):
        return {
            (0, 18200.0, 0.0),
            (18201, 45000.0, 19.0),
            (45001, 120000.0, 32.5),
            (120001, 180000.0, 37.0),
            (180001, None, 45.0)
        }

    # staticmethod:: welcome_message (string)
    def initi_welcome_message(self):
        return '\n,   Welcome to the ERP System!,,,'

    # staticmethod:: served_countries (Tuple)
    def served_countries(self):
        return ({
                    "Country": "Australia",
                    "Country Code": "AUS",
                    "Curr Code": "AUD",
                    "Rate to AUD": 1.0,
                },
                {
                    "Country": "HongKong",
                    "Country Code": "HKG",
                    "Curr Code": "HKD",
                    "Rate to AUD": 5.6,
                },
                {
                    "Country": "China",
                    "Country Code": "CHN",
                    "Curr Code": "CNY",
                    "Rate to AUD": 4.6,
                },
                {
                    "Country": "Malaysia",
                    "Country Code": "MYS",
                    "Curr Code": "MYR",
                    "Rate to AUD": 3.2,
                },
                {
                    "Country": "Vietnam",
                    "Country Code": "VNM",
                    "Curr Code": "VND",
                    "Rate to AUD": 16000.0,
                },
                {
                    "Country": "India",
                    "Country Code": "IND",
                    "Curr Code": "INR",
                    "Rate to AUD": 55.3,
                },)

    # staticmethod:: postcode_format (Set)
    def served_postcode_format(self):
        return {
            ("Australia", r'^\d{4}$'),
            ("Hong Kong", r'^0000$'),
            ("China", r'^\d{6}$'),
            ("Malaysia", r'^\d{5}$'),
            ("Vietnam", r'^\d{6}$'),
            ("India", r'^\d{6}$'),
        }

    # Function gen_customer_headers: Gen customer customize headers
    def gen_customer_headers(self):
        headers = list(self.enterprise_data.col_widths['customers'].keys())
        if 'postcode' in headers:
            headers[headers.index('postcode')] = 'pc'
        headers.insert(0, "No.")
        return headers

    # Function gen_product_headers: Gen product customize headers
    def gen_product_headers(self):
        headers = list(self.enterprise_data.col_widths['products'].keys())
        headers[headers.index('product_id')] = 'pid'
        headers[headers.index('stock_quantity')] = 'qty'
        return headers

    # Function gen_product_headers: Gen product customize headers
    def gen_order_headers(self):
        headers = list(self.enterprise_data.col_widths['orders'].keys())
        return headers


""" Part 5 - Main Class """

# Main application class for the ERP system
class ERPSystemApp:
    # Constructor
    def __init__(self, for_test_mode=False):
        # Create Utility Classes/Methods
        self.check_valid_method = CheckValidator()
        self.transform_data_method = DataTransformer()

        # Init the controllers/managers
        self.product_manager = ProductController()
        self.order_manager = OrderController(self.product_manager)
        self.customer_manager = CustomerController(self.order_manager)

        # Load the datasets into the controllers/managers
        self.enterprise_data = EnterpriseData()
        self.erp_data = ERPSystemData()
        self.load_data()

        # Init the ERP UI for the startup
        if not for_test_mode:
            self.welcome_message()
            self.menu_page()

    """ Part 5.1 - System Interaction Functions """

    def load_data(self):
        # Load customers
        for customer in self.enterprise_data.customers:
            try:
                self.customer_manager.create_customer(
                    customer.first_name,
                    customer.last_name,
                    customer.dob,
                    customer.email,
                    customer.phone,
                    customer.country,
                    customer.city,
                    customer.postcode,
                    customer.created_at,
                    customer.updated_at,
                )
            except Exception as er_msg:
                print(f"Error: loading customer '{customer.email}' - {str(er_msg)}")

        # Load products
        for prd in self.enterprise_data.products:
            try:
                self.product_manager.create_product(
                    prd.product_id,
                    prd.product_name,
                    prd.price,
                    prd.category,
                    prd.stock_quantity,
                    prd.created_at,
                    prd.updated_at,
                )
            except Exception as er_msg:
                print(f"Error: loading product '{prd.product_id}' - {str(er_msg)}")

        # Load orders
        for order in self.enterprise_data.orders:
            try:
                self.order_manager.create_order(
                    order.order_id,
                    order.order_date,
                    order.total_price,
                    [op.__dict__ for op in order.order_products],
                    order.payment_method,
                    order.order_status,
                    order.created_at,
                    order.updated_at,
                    order.customer_email,  # Added customer_email
                )
            except Exception as er_msg:
                print(f"Error: loading order '{order.order_id}' - {str(er_msg)}")

    # Function welcome_message: Prompt welcome message and allow dirty strings
    def welcome_message(self):
        print(self.erp_data.welcome_message.strip('\n, '))

    # Function get_cur_location: Get the current location on the console
    def get_cur_location(self):
        return os.getcwd()

    # Function transform_order_date: Transform order date
    def transform_order_date(self, order):
        return datetime.strptime(order.order_date, '%d-%m-%Y')

    # Function get_sales_quantity: Get sales quantity
    def get_sales_quantity(self, item):
        return item[1]

    # Main menu function to provide options
    def menu_page(self):
        menu_options = {
            '1': ("Create a new Customer", self.create_new_customer),
            '2': ("Create Payslip", self.create_payslip),  # New functionality for creating payslip
            '3': ("View Customer", self.display_customers),
            '4': ("View Currency Conversion Table Data", self.view_currency_conversion_table_data),  # New functionality for currency data
            '5': ("View Employee Information", self.view_employee_info),  # New functionality for employee info
            '6': ("View All Payslips", self.view_all_payslips),  # New functionality for viewing all payslips
            '7': ("View Victoria Tax Table Data", self.view_victoria_tax_table_data),  # New functionality for tax data
            '8': ("View Product", self.display_product_table),
            '9': ("View Order", self.display_orders),
            '10': ("View Chart", self.display_charts),
            '11': ("Export Charts (.png)", self.export_all_charts),
            '12': ("Export Customer (customers.csv)", lambda: self.export_csv('customer')),
            '13': ("Export Product (products.csv)", lambda: self.export_csv('product')),
            '14': ("Export Orders (orders.json)", self.export_orders_json),
            '15': ("Export Report (docx)", self.export_report),
            '16': ("Import Customer (customers.csv)", lambda: self.read_csv('customer')),
            '17': ("Import Product (products.csv)", lambda: self.read_csv('product')),
            '18': ("Import Order (orders.json)", lambda: self.read_json('order')),
        }

        while True:
            print(f"""
            --- Main Menu ---
    
            [ Create ]
            1. {menu_options['1'][0]}
            2. {menu_options['2'][0]}
            
            [ Read ] [ Update ] [ Delete ]
            3. {menu_options['3'][0]}
            
            [ Read ] [ Update ]
            4. {menu_options['4'][0]}
            
            [ Read ] [ Delete ]
            5. {menu_options['5'][0]}
            6. {menu_options['6'][0]}
            
            [ Read ]
            7. {menu_options['7'][0]}
            8. {menu_options['8'][0]}
            9. {menu_options['9'][0]}
            10. {menu_options['10'][0]}
    
            [ Export ]
            11. {menu_options['11'][0]}
            12. {menu_options['12'][0]}
            13. {menu_options['13'][0]}
            14. {menu_options['14'][0]}
            15. {menu_options['15'][0]}
    
            [ Import ]
            Please note that your files must be in the below path:
            {self.get_cur_location()}
            16. {menu_options['16'][0]}
            17. {menu_options['17'][0]}
            18. {menu_options['18'][0]}
            """)

            usr_chs = input("Enter your choice: ")

            # Check if the user's choice is valid
            if usr_chs in menu_options:
                # Execute the corresponding method
                menu_options[usr_chs][1]()
            else:
                print("Invalid choice. Please try again.")



    """ Part 5.2 - Create Functions """

    # Create Function create_new_customer: New Customer
    def create_new_customer(self):
        print("\nCreate a New Customer")
        email = input("Enter email: ").strip()

        # Check
        if not self.check_valid_method.is_email(email) or self.customer_manager.get_customer_by_email(email):
            print("Invalid email or customer already exists.")
            return
        first_name = input("Enter first name: ")
        if not self.check_valid_method.is_alpha(first_name): return
        last_name = input("Enter last name: ")
        if not self.check_valid_method.is_alpha(last_name): return
        dob = input("Enter date of birth (YYYY-MM-DD): ").strip()
        if not self.check_valid_method.is_valid_date_of_birth(dob): return
        phone = input("Enter phone number (0000-000-000): ").strip()
        if not self.check_valid_method.is_phone_num(phone): return
        country = input("Enter country: ").strip()
        if not self.check_valid_method.is_served_country(country, self.erp_data.full_served_countries): return
        city = input("Enter city: ")
        if not self.check_valid_method.is_alpha(city): return
        postcode = input("Enter postcode: ").strip()
        if not self.check_valid_method.is_postcode(country, postcode, self.erp_data.full_served_countries): return

        # Confirm
        print(
            f"\nConfirm customer info:\nEmail: {email}\nFirst Name: {first_name}\nLast Name: {last_name}\nDOB: {dob}"
            f"\nPhone: {phone}\nCountry: {country}\nCity: {city}\nPostcode: {postcode}")
        if input("Is this information correct? (y/n): ").strip().lower() != 'y':
            print("Customer creation cancelled.")
            return

        created_at = updated_at = datetime.now().strftime(
            "%Y-%m-%d %H:%M:%S"
        )

        # Try Create
        try:
            self.customer_manager.create_customer(first_name, last_name, dob, email, phone, country, city, postcode,
                                                  created_at, updated_at)
            print("Customer created successfully.")
        except Exception as er_msg:
            print("Error: creating customer:", str(er_msg))

    """ Part 5.3 - Read Functions """

    # Read Function display_table: View aligned tables
    def display_table(self, t_data, t_headers, t_col_widths):
        format_of_headers = ''
        for a_header in t_headers:
            format_of_headers += self.transform_data_method.pad_string(a_header, t_col_widths.get(a_header, 10) + 2)
        print(format_of_headers)
        print('-' * len(format_of_headers))
        for a_row in t_data:
            format_of_row = ''
            for a_header in t_headers:
                t_value = a_row.get(a_header, "N/A")
                format_of_row += self.transform_data_method.pad_string(t_value, t_col_widths.get(a_header, 10) + 2)
            print(format_of_row)

    # Read Function display_product_table: Display Products table for more details
    def display_product_table(self):
        products = self.product_manager.get_all_products()
        if not products:
            print("No products available.")
            return

        # Customize the Header
        col_widths = self.enterprise_data.col_widths['products']
        col_widths_with_no = {
            'No.': 4,
            'pid': col_widths['product_id'],
            'product_name': col_widths['product_name'],
            'price': col_widths['price'],
            'category': col_widths['category'],
            'qty': col_widths['stock_quantity'],
            'created_at': col_widths['created_at'],
            'updated_at': col_widths['updated_at']
        }
        headers = self.erp_data.custom_headers['products']

        data = [
            {'No.': indx, 'pid': p.product_id, 'product_name': p.product_name, 'price': str(p.price),
             'category': p.category, 'qty': str(p.stock_quantity),
             'created_at': self.transform_data_method.to_std_datetimeformat(p.created_at),
             'updated_at': self.transform_data_method.to_std_datetimeformat(p.updated_at)}
            for indx, p in enumerate(products, start=1)
        ]

        self.display_table(data, headers, col_widths_with_no)

        while True:
            usr_chs = input("\nEnter product number to view details, or 'b' to go back: ").strip().lower()
            if usr_chs == 'b': break
            if self.check_valid_method.is_numeric(usr_chs) and 1 <= int(usr_chs) <= len(products):
                self.display_product_details(products[int(usr_chs) - 1])
            else:
                print("Invalid choice. Please try again.")

    # Read Function display_product_details: View Product's details
    def display_product_details(self, product):
        print(f"""
            Product Details for {product.product_name}:
            Product ID: {product.product_id}
            Name: {product.product_name}
            Price: {product.price}
            Category: {product.category}
            Quantity (qty): {product.stock_quantity}
            Created Date: {self.transform_data_method.to_std_datetimeformat(product.created_at)}
            Updated Date: {self.transform_data_method.to_std_datetimeformat(product.updated_at)}""")

    # Read Function display_customers: Display Customers for more details
    def display_customers(self):
        while True:
            customers = self.customer_manager.get_all_customers()
            if not customers:
                print("No customers available.")
                return

            self.display_customer_table(customers)

            usr_chs = input(
                "\nEnter customer number to view details, or 'b' to go back: "
            ).strip()
            if usr_chs.lower() == 'b':
                break
            elif self.check_valid_method.is_numeric(usr_chs):
                idx = int(usr_chs) - 1
                if 0 <= idx < len(customers):
                    selected_customer = customers[idx]
                    self.display_customer_details(selected_customer)
                else:
                    print("Invalid customer number. Please try again.")
            else:
                print("Invalid choice. Please try again.")

    # Read Function display_customer_table: Display Customers table by custom
    def display_customer_table(self, customers):
        col_widths = self.enterprise_data.col_widths['customers']
        if 'postcode' in col_widths:
            col_widths['pc'] = col_widths.pop('postcode')

        headers = self.erp_data.custom_headers['customers']

        col_widths_with_no = {'No.': 4}
        col_widths_with_no.update(col_widths)

        data = [
            {
                'No.': indx,
                'first_name': a_cus.first_name,
                'last_name': a_cus.last_name,
                'dob': a_cus.dob,
                'email': a_cus.email,
                'phone': a_cus.phone,
                'country': a_cus.country,
                'city': a_cus.city,
                'pc': a_cus.postcode,
                'created_at': self.transform_data_method.to_std_datetimeformat(a_cus.created_at),
                'updated_at': self.transform_data_method.to_std_datetimeformat(a_cus.updated_at)
            }
            for indx, a_cus in enumerate(customers, start=1)
        ]

        self.display_table(data, headers, col_widths_with_no)

    # Read Function display_customer_details: View each Customer's details
    def display_customer_details(self, customer):
        while True:
            print(f"""
            Customer Details:
            {customer.first_name} {customer.last_name}
            Email: {customer.email}
            Date of Birth: {customer.dob}
            Phone: {customer.phone}
            Country: {customer.country}
            City: {customer.city}
            Postcode: {customer.postcode}
            Created Date: {self.transform_data_method.to_std_datetimeformat(customer.created_at)}
            Updated Date: {self.transform_data_method.to_std_datetimeformat(customer.updated_at)}""")

            print("\nOptions:")
            print("1. Update Customer")
            print("2. Delete Customer")
            print("b. Go Back")
            usr_chs = input("Enter your choice: ").strip().lower()
            if usr_chs == '1':
                self.update_customer(customer)
            elif usr_chs == '2':
                self.delete_customer(customer)
                break
            elif usr_chs == 'b':
                break
            else:
                print("Invalid choice. Please try again.")

    # Read Function display_charts: View a chart by selection
    def display_charts(self):
        while True:
            print(f"""
            Reports Menu:
            1. Top-Selling Products
            2. Customer Location Distribution
            3. Customer Age Distribution
            b. Go Back\n""")

            usr_chs = input("Enter your choice: ")
            if usr_chs == '1':
                self.gen_bar_chart_top_selling_products()
            elif usr_chs == '2':
                self.gen_customer_location_distribution_chart()
            elif usr_chs == '3':
                self.gen_customer_age_distribution_chart()
            elif usr_chs == 'b':
                break
            else:
                print("Invalid choice. Please try again.")

    # Read Function display_orders: Display Orders for more details
    def display_orders(self):
        orders = self.order_manager.get_all_orders()
        if not orders:
            print("No orders available.")
            return
        for indx, a_order in enumerate(orders, start=1):
            print(
                f"{indx}. Order ID: {a_order.order_id}, Customer Email: {a_order.customer_email}, Total Price: {a_order.total_price}")

        while True:
            usr_chs = input("\nEnter order number to view details, or 'b' to go back: ").strip().lower()
            if usr_chs == 'b': break
            if self.check_valid_method.is_numeric(usr_chs) and 1 <= int(usr_chs) <= len(orders):
                self.display_order_details(orders[int(usr_chs) - 1])
            else:
                print("Invalid choice. Please try again.")

    # Read Function display_order_details: View Order's details
    def display_order_details(self, order):
        print(f"""
            Order Details for Order ID: {order.order_id}
            Order Date: {order.order_date}
            Customer Email: {order.customer_email}
            Order Status: {order.order_status}
            Payment Method: {order.payment_method}
            Total Price: {order.total_price}
            Created Date: {self.transform_data_method.to_std_datetimeformat(order.created_at)}
            Updated Date: {self.transform_data_method.to_std_datetimeformat(order.updated_at)}""")
        print("\nOrdered Products:")
        for op in order.order_products:
            print(
                f"- Product ID: {op.product_id}, Name: {op.product_name}, "
                f"Quantity: {op.quantity}, Price per Unit: {op.price_per_unit}, Total Price: {op.total_price}"
            )

    """ Part 5.4 - Update Functions """

    # Update Function update_customer: Update a Customer details
    def update_customer(self, customer):
        print("\nUpdate Customer")
        # Check Input one by one
        n_first_name = input(
            f"Enter new first name (current: {customer.first_name}): "
        ).strip()
        if n_first_name:
            if not self.check_valid_method.is_alpha(n_first_name):
                print("Invalid first name. Update cancelled.")
                return
            customer.first_name = n_first_name
        n_last_name = input(
            f"Enter new last name (current: {customer.last_name}): "
        ).strip()
        if n_last_name:
            if not self.check_valid_method.is_alpha(n_last_name):
                print("Invalid last name. Update cancelled.")
                return
            customer.last_name = n_last_name
        n_dob = input(
            f"Enter new date of birth (YYYY-MM-DD) (current: {customer.dob}): "
        ).strip()
        if n_dob:
            if not self.check_valid_method.is_valid_date_of_birth(n_dob):
                print("Invalid date of birth. Update cancelled.")
                return
            customer.dob = n_dob
        n_phone = input(
            f"Enter new phone number (current: {customer.phone}): "
        ).strip()
        if n_phone:
            if not self.check_valid_method.is_phone_num(n_phone):
                print("Invalid phone number. Update cancelled.")
                return
            customer.phone = n_phone
        n_country = input(
            f"Enter new country (current: {customer.country}): "
        ).strip()
        if n_country:
            if not self.check_valid_method.is_served_country(
                    n_country, self.erp_data.full_served_countries
            ):
                print("Invalid country name. Update cancelled.")
                return
            customer.country = n_country

        n_city = input(f"Enter new city (current: {customer.city}): ").strip()
        if n_city:
            if not self.check_valid_method.is_alpha(n_city):
                print("Invalid city name. Update cancelled.")
                return
            customer.city = n_city
        n_postcode = input(
            f"Enter new postcode (current: {customer.postcode}): "
        ).strip()
        if n_postcode:
            if not self.check_valid_method.is_postcode(
                    n_country, n_postcode, self.erp_data.full_served_countries
            ):
                print("Invalid postcode. Update cancelled.")
                return
            customer.postcode = n_postcode
        # Update the updated_at as current
        customer.updated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        print("Customer updated successfully.")

    # Update Function update_customer_by_email: Run update_customer b4 checking by Email
    def update_customer_by_email(self):
        input_email = input("Enter the email of the customer to update: ").strip()
        the_customer = self.customer_manager.get_customer_by_email(input_email)
        if not the_customer:
            print("Customer not found.")
            return
        self.update_customer(the_customer)

    """ Part 5.5 - Delete Functions """

    # Delete Function delete_customer: Delete a customer by passing a Customer object
    def delete_customer(self, customer):
        while True:
            confirm = input(
                f"Are you sure you want to delete customer '{customer.first_name} {customer.last_name}'? (y/n): "
            ).strip().lower()
            if confirm == 'y':
                self.order_manager.delete_orders_by_customer_email(customer.email)
                self.customer_manager.delete_orders_by_customer(customer.email)
                print("Customer and related orders deleted successfully.")
                break
            else:
                print("Deletion cancelled.")
                break

    """ Part 5.6 - Export Functions """

    # Export Function export_report: Create a Report, then export
    def export_report(self):
        report_contents = self.create_report_content()
        set_filename = f"marketing_report_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        ms_word_doc = Document()
        self.create_report_document(ms_word_doc, report_contents)
        ms_word_doc.save(set_filename)
        print(f'Report "{set_filename}" created successfully. \nLocation: "{self.get_cur_location()}\\{set_filename}"')

    # Export Function export_csv: Create a csv by data list, then export
    def export_csv(self, object_name):
        name_in_col_widths = object_name + 's'
        assume_filename = "".join([name_in_col_widths, '.csv'])
        try:
            headers = self.erp_data.custom_headers[object_name + 's']
            csv_data = []
            if object_name == 'customer':
                if "No." not in headers:
                    headers.insert(0, "No.")
                customers = self.customer_manager.get_all_customers()
                if not customers:
                    print(f"No customers found to export to {assume_filename}")
                    return
                csv_data = [
                    {
                        "No.": idx + 1,
                        "first_name": customer.first_name,
                        "last_name": customer.last_name,
                        "dob": customer.dob,
                        "email": customer.email,
                        "phone": customer.phone,
                        "country": customer.country,
                        "city": customer.city,
                        "pc": customer.postcode,
                        "created_at": customer.created_at,
                        "updated_at": customer.updated_at,
                    }
                    for idx, customer in enumerate(customers)
                ]
            elif object_name == 'product':
                products = self.product_manager.get_all_products()
                if not products:
                    print(f"No products found to export to {assume_filename}")
                    return
                csv_data = [
                    {
                        "pid": product.product_id,
                        "product_name": product.product_name,
                        "price": product.price,
                        "category": product.category,
                        "qty": product.stock_quantity,
                        "created_at": product.created_at,
                        "updated_at": product.updated_at,
                    }
                    for product in self.product_manager.get_all_products()
                ]
            # Write the data to the CSV file
            with open(assume_filename, 'w', newline='') as file:
                csv_writer = csv.DictWriter(file, fieldnames=headers)
                csv_writer.writeheader()
                csv_writer.writerows(csv_data)
                csv_file_message = f'CSV file "{assume_filename}" has been created successfully.'
                file_loca_message = f'Location: "{self.get_cur_location()}\\{assume_filename}"'
                print("{} \n {}".format(csv_file_message, file_loca_message))
        except Exception as er_msg:
            print(f"Error: exporting {object_name} list - {str(er_msg)}")

    # Export Function export_orders_json: Create a jason by Orders, then export
    def export_orders_json(self):
        set_filename = 'orders.json'
        try:
            orders_data = []
            orders = self.order_manager.get_all_orders()
            if not orders:
                print(f"No orders found to export to {set_filename}")
                return
            for order in orders:
                order_dict = vars(order)
                order_dict['order_products'] = [
                    vars(prd) for prd in order.order_products
                ]
                orders_data.append(order_dict)
            with open(set_filename, 'w') as json_file:
                json.dump(orders_data, json_file, indent=4)
            print(
                f'JSON file "{set_filename}" created successfully. \nLocation: "{self.get_cur_location()}\\{set_filename}"')
        except Exception as er_msg:
            print(f"Error: exporting orders: {er_msg}")

    # Export Function export_all_charts: Call all the gen chart functions as export
    def export_all_charts(self):
        print("\nExporting all charts as .png files...")
        self.gen_bar_chart_top_selling_products(export=True)
        self.gen_customer_location_distribution_chart(export=True)
        self.gen_customer_age_distribution_chart(export=True)
        file_loca_message = f'Location: "{self.get_cur_location()}\\"'
        print("All charts have been exported successfully.")
        print(file_loca_message)

    """ Part 5.7 - Import Functions """

    # Import Function read_csv: Read the csv and save the data by inputted object_name name
    def read_csv(self, object_name):
        assume_filename = "".join([object_name + 's', '.csv'])
        try:
            with open(assume_filename, newline='') as csv_file:
                if object_name == 'customer':
                    correct_headers = self.erp_data.custom_headers[object_name + 's']
                    csv_reader = csv.DictReader(csv_file, fieldnames=correct_headers)
                    # Skips the heading - Using next() method
                    next(csv_file)
                    for row in csv_reader:
                        try:
                            # Check
                            first_name = row['first_name'].strip()
                            if not self.check_valid_method.is_alpha(first_name):
                                print(f"Invalid first name '{first_name}'. Skipping this record.")
                                continue
                            last_name = row['last_name'].strip()
                            if not self.check_valid_method.is_alpha(last_name):
                                print(f"Invalid last name '{last_name}'. Skipping this record.")
                                continue
                            dob = row['dob'].strip()
                            dob = self.transform_data_method.to_std_dateformat(dob)
                            if not self.check_valid_method.is_valid_date_of_birth(dob):
                                print(f"Invalid date of birth '{dob}'. Skipping this record.")
                                continue
                            email = row['email'].strip()
                            if not self.check_valid_method.is_email(email):
                                print(f"Invalid email '{email}'. Skipping this record.")
                                continue
                            phone = row['phone'].strip()
                            if not self.check_valid_method.is_phone_num(phone):
                                print(f"Invalid phone number '{phone}'. Skipping this record.")
                                continue
                            country = row['country'].strip()
                            if not self.check_valid_method.is_served_country(country,
                                                                             self.erp_data.full_served_countries):
                                print(f"Invalid country '{country}'. Skipping this record.")
                                continue
                            city = row['city'].strip()
                            if not self.check_valid_method.is_alpha(city):
                                print(f"Invalid city '{city}'. Skipping this record.")
                                continue
                            postcode = row['pc'].strip()
                            if not self.check_valid_method.is_postcode(country, postcode,
                                                                       self.erp_data.full_served_countries):
                                print(f"Invalid postcode '{postcode}'. Skipping this record.")
                                continue
                            created_at = row['created_at'].strip()
                            created_at = self.transform_data_method.to_std_datetimeformat(created_at)
                            if not self.check_valid_method.is_datetime(created_at):
                                print(f"Invalid created datetime '{created_at}'. Skipping this record.")
                                continue
                            updated_at = row['updated_at'].strip()
                            updated_at = self.transform_data_method.to_std_datetimeformat(updated_at)
                            if not self.check_valid_method.is_datetime(updated_at):
                                print(f"Invalid updated datetime '{updated_at}'. Skipping this record.")
                                continue
                            # Add or Edit
                            self.customer_manager.create_customer(
                                row['first_name'], row['last_name'], row['dob'], row['email'], row['phone'],
                                row['country'], row['city'], row['pc'], row['created_at'], row['updated_at']
                            )
                        except Exception as er_msg:
                            print(f"Error: adding customer '{row['email']}' - {str(er_msg)}")

                elif object_name == 'product':
                    csv_reader = csv.DictReader(csv_file)
                    for row in csv_reader:
                        try:
                            # Check
                            product_id = row['pid'].strip()
                            if not product_id:
                                print("Product ID is missing. Skipping this record.")
                                continue
                            product_name = row['product_name'].strip()
                            if not product_name:
                                print(f"Product name is missing for ID '{product_id}'. Skipping this record.")
                                continue
                            price = row['price'].strip()
                            if not self.check_valid_method.is_decimal(price):
                                print(f"Invalid price '{price}' for product '{product_name}'. Skipping this record.")
                                continue
                            price = float(price)
                            category = row['category'].strip()
                            if not category:
                                print(f"Category is missing for product '{product_name}'. Skipping this record.")
                                continue
                            stock_quantity = row['qty'].strip()
                            if not self.check_valid_method.is_numeric(stock_quantity):
                                print(
                                    f"Invalid stock quantity '{stock_quantity}' for product '{product_name}'. Skipping this record.")
                                continue
                            stock_quantity = int(stock_quantity)
                            created_at = row['created_at'].strip()
                            created_at = self.transform_data_method.to_std_datetimeformat(created_at)
                            if not self.check_valid_method.is_datetime(created_at):
                                print(f"Invalid created datetime '{created_at}'. Skipping this record.")
                                continue
                            updated_at = row['updated_at'].strip()
                            updated_at = self.transform_data_method.to_std_datetimeformat(updated_at)
                            if not self.check_valid_method.is_datetime(updated_at):
                                print(f"Invalid updated datetime '{updated_at}'. Skipping this record.")
                                continue
                            # Add
                            self.product_manager.create_product(
                                product_id, product_name, price, category, stock_quantity,
                                created_at, updated_at
                            )
                        except Exception as er_msg:
                            print(f"Error: adding product '{product_name}' - {str(er_msg)}")

                file_loca_message = f'Location: "{self.get_cur_location()}\\{assume_filename}"'
                csv_file_message = f'CSV file "{assume_filename}" has been imported successfully.'
                print("{} \n {}".format(file_loca_message, csv_file_message))
        except Exception as er_msg:
            print(f"Error: importing {object_name} list - {str(er_msg)}")

    # Import Function read_json: Read the JSON and save the data by inputted object_name name
    def read_json(self, object_name):
        assume_filename = "".join([object_name + 's', '.json'])
        try:
            with open(assume_filename, 'r') as json_file:
                data = json.load(json_file)

                if object_name == 'order':
                    for order_data in data:
                        try:
                            # Check
                            for ordered_product in order_data['order_products']:
                                product = self.product_manager.get_product_by_id(ordered_product['product_id'])
                                if not product:
                                    print(
                                        f"Product with ID {ordered_product['product_id']} "
                                        f"does not exist. Skipping this product.")
                                    continue
                                required_product_fields = ['product_id', 'product_name', 'quantity', 'price_per_unit',
                                                           'total_price']
                                for field in required_product_fields:
                                    if field not in ordered_product:
                                        if field == 'price_per_unit':
                                            ordered_product['price_per_unit'] = product.price
                                        elif field == 'total_price':
                                            ordered_product['total_price'] = ordered_product['quantity'] * \
                                                                             ordered_product[
                                                                                 'price_per_unit']
                                        else:
                                            print(
                                                f"Error: Missing required field '{field}' in product "
                                                f"{ordered_product.get('product_name', '')}. Skipping this product.")
                                            continue

                            # Add or Edit
                            self.order_manager.create_order(
                                order_data['order_id'], order_data['order_date'], order_data['total_price'],
                                order_data['order_products'], order_data['payment_method'],
                                order_data['order_status'], order_data['created_at'],
                                order_data['updated_at'], order_data['customer_email']
                            )
                        except Exception as er_msg:
                            print(f"Error: adding order '{order_data['order_id']}' - {str(er_msg)}")
            print(f'JSON file "{assume_filename}" has been imported successfully.')
        except Exception as er_msg:
            print(f"Error: importing {object_name} list - {str(er_msg)}")

    """ Part 5.8 - Charts Functions """

    # Chart Function save_chart: Save the chart by its title
    def save_chart(self, fig, inpt_chart_title):
        set_filename = f"{inpt_chart_title.replace(' ', '_').lower()}.png"
        fig.savefig(set_filename)
        chart_file_message = f'Chart "{set_filename}" has been created successfully.'
        print(chart_file_message)
        plt.close(fig)

    # Chart Function gen_chart: Gen a chart
    def gen_chart(self, data, labels, chart_title, xlabel, ylabel, export=False):
        fig, ax = plt.subplots()
        ax.bar(labels, data)
        ax.set_title(chart_title)
        ax.set_xlabel(xlabel)
        ax.set_ylabel(ylabel)
        plt.xticks(rotation=45)
        plt.tight_layout()
        if export:
            self.save_chart(fig, chart_title)
        else:
            print("Please close your chart's window to continue.")
            plt.show()

        return fig, ax

    # Chart Function gen_bar_chart_top_selling_products: gen chart
    def gen_bar_chart_top_selling_products(self, export=False):
        pd_sales = {}
        for a_order in self.order_manager.get_all_orders():
            for or_pd in a_order.order_products:
                if isinstance(or_pd, dict):
                    pd_sales[or_pd['product_name']] = pd_sales.get(or_pd['product_name'], 0) + int(or_pd['quantity'])
                elif hasattr(or_pd, 'product_name') and hasattr(or_pd, 'quantity'):
                    pd_sales[or_pd.product_name] = pd_sales.get(or_pd.product_name, 0) + int(or_pd.quantity)
        if not pd_sales:
            print("No sales data available.")
            return
        self.gen_chart(
            list(pd_sales.values()),
            list(pd_sales.keys()),
            'Top-Selling Products',
            'Products',
            'Units Sold',
            export
        )

    # Chart Function gen_customer_location_distribution_chart: gen chart
    def gen_customer_location_distribution_chart(self, export=False):
        city_counts = {}
        for a_customer in self.customer_manager.get_all_customers():
            city_counts[a_customer.city] = city_counts.get(a_customer.city, 0) + 1
        if not city_counts:
            print("No customer location data available.")
            return
        self.gen_chart(
            list(city_counts.values()),
            list(city_counts.keys()),
            'Customer Location Distribution',
            'City',
            'Number of Customers',
            export
        )

    # Chart Function gen_customer_age_distribution_chart: gen chart
    def gen_customer_age_distribution_chart(self, export=False):
        customers = self.customer_manager.get_all_customers()
        if not customers:
            print("No customer data available.")
            return

        ages = [
            (datetime.now() - datetime.strptime(customer.dob, '%Y-%m-%d')).days // 365
            for customer in customers if customer.dob
        ]

        if not ages:
            print("No valid date of birth data available.")
            return

        # Create bins/ranges of ages by numpy
        range_of_ages = np.histogram(ages, bins=10)
        range_of_ages_labels = [f"{int(range_of_ages[1][i])}-{int(range_of_ages[1][i + 1])}" for i in
                                range(len(range_of_ages[1]) - 1)]

        self.gen_chart(
            range_of_ages[0],
            range_of_ages_labels,
            'Customer Age Distribution',
            'Age',
            'Number of Customers',
            export
        )

    """ Part 5.9 - Report Functions """

    # Report Function create_report_content: gen report's content
    def create_report_content(self):
        return {
            'kpis': self.gen_kpis_section(),
            'product_performance': self.gen_product_performance_section(),
            'order_customer_insights': self.gen_order_customer_insights_section(),
            'charts': self.get_chart_filenames(),
        }

    # Report Function gen_kpis_section: gen KPI's section
    def gen_kpis_section(self):
        tt_sales = sum(float(order.total_price) for order in self.order_manager.get_all_orders())
        tt_customers = len(self.customer_manager.get_all_customers())
        tt_orders = len(self.order_manager.get_all_orders())
        avg_order_value = tt_sales / tt_orders if tt_orders > 0 else 0
        customer_retention_rate = 100  # Placeholder logic

        return {
            'Total Sales': f"${tt_sales:,.2f}",
            'Total Customers': tt_customers,
            'Average Order Value': f"${avg_order_value:,.2f}",
            'Customer Retention Rate': f"{customer_retention_rate:.2f}%"  # Simplified
        }

    # Report Function gen_product_performance_section: gen product's performance section
    def gen_product_performance_section(self):
        pd_sales = {}
        for a_order in self.order_manager.get_all_orders():
            for op in a_order.order_products:
                if isinstance(op, dict):
                    pd_sales[op['product_name']] = pd_sales.get(op['product_name'], 0) + int(op['quantity'])
                elif hasattr(op, 'product_name') and hasattr(op, 'quantity'):
                    pd_sales[op.product_name] = pd_sales.get(op.product_name, 0) + int(op.quantity)
        best_sale = max(pd_sales.items(), key=self.get_sales_quantity, default=("N/A",))[0]
        the_low_stock = [f"{p.product_name} – {p.stock_quantity} units" for p in self.product_manager.get_all_products()
                         if p.stock_quantity < 10]

        return {
            'Best-Selling Product': best_sale,
            'Low Stock Alert': the_low_stock,
            'Product Table': [
                {'Product Name': p.product_name, 'Units Sold': pd_sales.get(p.product_name, 0),
                 'Total Revenue': f"${pd_sales.get(p.product_name, 0) * p.price:,.2f}"}
                for p in self.product_manager.get_all_products()
            ]
        }

    # Report Function gen_order_customer_insights_section: gen customer's insights section
    def gen_order_customer_insights_section(self):
        rc_orders = sorted(self.order_manager.get_all_orders(), key=self.transform_order_date, reverse=True)[:5]
        cities_of_customer = {}
        for a_customer in self.customer_manager.get_all_customers():
            cities_of_customer[a_customer.city] = cities_of_customer.get(a_customer.city, 0) + 1

        return {
            'Recent Orders': [{'Order ID': o.order_id, 'Date': o.order_date, 'Total': f"${float(o.total_price):,.2f}"}
                              for o in rc_orders],
            'Customer Demographics': [{'City': city, 'Number of Customers': count} for city, count in
                                      cities_of_customer.items()]
        }

    # Report Function get_chart_filenames: get charts filename
    def get_chart_filenames(self):
        return [
            'top-selling_products.png',
            'payment_methods_distribution.png',
            'customer_location_distribution.png',
            'customer_age_distribution.png'
        ]

    # Report Function create_report_document: gen report's Microsoft doc
    def create_report_document(self, ms_wd_doc, contents):
        # Add Title
        ms_wd_doc.add_heading('Marketing Report', 0)

        # 1st Section – KPIs
        ms_wd_doc.add_heading('1. KPIs', level=1)
        for key, value in contents['kpis'].items():
            ms_wd_doc.add_paragraph(f"• {key}: {value}")

        # 2nd Section – Product Performance
        ms_wd_doc.add_heading('2. Product Performance', level=1)
        ms_wd_doc.add_paragraph(f"• Best-Selling Product: {contents['product_performance']['Best-Selling Product']}")
        ms_wd_doc.add_paragraph(f"• Low Stock Alert:")
        for alert in contents['product_performance']['Low Stock Alert']:
            ms_wd_doc.add_paragraph(f"  - {alert}", style='List Bullet')

        # 3rd Section – Product Table
        ms_wd_doc.add_heading('3. Product Table', level=1)
        table = ms_wd_doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Product Name'
        hdr_cells[1].text = 'Units Sold'
        hdr_cells[2].text = 'Total Revenue'
        for a_item in contents['product_performance']['Product Table']:
            row_cells = table.add_row().cells
            row_cells[0].text = a_item['Product Name']
            row_cells[1].text = str(a_item['Units Sold'])
            row_cells[2].text = a_item['Total Revenue']

        # 4th Section – Order and Customer Insights
        ms_wd_doc.add_heading('4. Order and Customer Insights', level=1)
        table = ms_wd_doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Order ID'
        hdr_cells[1].text = 'Date'
        hdr_cells[2].text = 'Total'
        for a_item in contents['order_customer_insights']['Recent Orders']:
            row_cells = table.add_row().cells
            row_cells[0].text = a_item['Order ID']
            row_cells[1].text = a_item['Date']
            row_cells[2].text = a_item['Total']

        # 5th - Customer Demographics Table
        ms_wd_doc.add_heading('5. Customer Demographics', level=1)
        table = ms_wd_doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'City'
        hdr_cells[1].text = 'Number of Customers'
        for a_item in contents['order_customer_insights']['Customer Demographics']:
            row_cells = table.add_row().cells
            row_cells[0].text = a_item['City']
            row_cells[1].text = str(a_item['Number of Customers'])

        # 6th - Charts
        ms_wd_doc.add_heading('6. Charts', level=1)
        self.export_all_charts()
        for a_chart_filename in contents['charts']:
            if os.path.exists(a_chart_filename):
                ms_wd_doc.add_picture(a_chart_filename, width=Inches(6))
                ms_wd_doc.add_paragraph(a_chart_filename.replace('_', ' ').replace('.png', '').title())
            else:
                ms_wd_doc.add_paragraph(f"Chart '{a_chart_filename}' not found.")

""" Part 6 - Start the program """
# Main Function
if __name__ == "__main__":
    ERPSystemApp()

